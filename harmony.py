# streamlit_app.py
import streamlit as st
import socket, sys, platform
import datetime
import subprocess
import uuid, io
from io import BytesIO
import re, hashlib
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from pydantic import BaseModel
import json, os, pathlib
import requests #for calling Google Custom Search API
import pandas as pd
from typing import Tuple
import traceback
from pathlib import Path
from fpdf import FPDF
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import google.generativeai as genai
import openai
import shutil
import base64
import time 
start_time = time.time() 
import threading, getpass
from io import BytesIO
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
pdfmetrics.registerFont(UnicodeCIDFont('HeiseiKakuGo-W5'))  # Japanese font but supports emoji
from matplotlib import font_manager as fm
import os
import builtins  # add at top of file
import pytesseract
import matplotlib.font_manager as fm
import sys
import io
import platform
import logging
from pathlib import Path
import tempfile
import gzip
import json
from pathlib import Path
import requests
import re
import smtplib
import mimetypes
from email.message import EmailMessage
import sqlite3
import bcrypt

FLASHMIND_KEY = st.secrets["FLASHMIND_KEY"]
OMNICORE_KEY = st.secrets["OMNICORE_KEY"]

DEFAULT_FROM = st.secrets["DEFAULT_FROM"]


def get_secret(name, default=""):

    try:
        return st.secrets[name]

    except:

        return default

# =========================================================
# FORCE SESSION RESET
# =========================================================

if "session_reset_done" not in st.session_state:

    for k in list(st.session_state.keys()):

        del st.session_state[k]

    st.session_state.session_reset_done = True



# =========================================================
# UI KEY PRIORITY SYSTEM
# UI key overrides .env
# =========================================================

def resolve_api_key(ui_key, default_key):

    if ui_key and str(ui_key).strip():
        return ui_key.strip()

    return default_key

# ===============================
# SAFE LOGGING SETUP (STREAMLIT + WINDOWS SAFE)
# ===============================

import logging
import platform
import tempfile
from pathlib import Path

import streamlit as st
from logging.handlers import RotatingFileHandler

# =========================================================
# AWS / STREAMLIT RUNTIME DEFAULTS
# =========================================================

BASE_DIR = Path(__file__).resolve().parent

DEFAULT_USER = "guest"

DEFAULT_ROLE = "viewer"

DEFAULT_ORGANIZATION = "Harmony"

# =========================================================
# SESSION DEFAULTS
# =========================================================

if "authenticated" not in st.session_state:

    st.session_state.authenticated = False

if "username" not in st.session_state:

    st.session_state.username = DEFAULT_USER

if "role" not in st.session_state:

    st.session_state.role = DEFAULT_ROLE

if "organization" not in st.session_state:

    st.session_state.organization = DEFAULT_ORGANIZATION

# =========================================================
# ACTIVE USER OBJECT
# =========================================================

USER = {

    "email": st.session_state.username,

    "identity": st.session_state.organization

}

ROLE = st.session_state.role




def read_readme():

    readme_path = BASE_DIR / "README.md"

    if readme_path.exists():
        return readme_path.read_text(
            encoding="utf-8",
            errors="ignore"
        )

    return "README not found."

# =========================================================
# HARMONY ROOT DIRECTORY
# Windows + Linux + AWS + Streamlit
# =========================================================


import platform
from pathlib import Path

# Detect server mode
IS_SERVER = os.environ.get("HARMONY_SERVER", "0") == "1"

def get_harmony_dir():
    """
    Resolve Harmony data directory safely
    (Windows / Linux / AWS / Streamlit / EXE)
    """

    custom_dir = os.environ.get("HARMONY_HOME")

    if custom_dir:
        return Path(custom_dir)

    # ======================================
    # AWS / VPS / SERVER MODE
    # ======================================
    if IS_SERVER:
        return Path("/home/ubuntu/Harmony-AI/.harmony_server")

    # ======================================
    # LOCAL DESKTOP MODE
    # ======================================
    return Path.home() / ".harmony"


# ======================================
# SERVER INFO
# ======================================

def get_server_info():

    try:
        hostname = socket.gethostname()
        local_ip = socket.gethostbyname(hostname)

    except Exception:
        hostname = "unknown"
        local_ip = "unknown"

    # AWS-safe public IP placeholder
    public_ip = "aws-managed"

    return {
        "hostname": hostname,
        "local_ip": local_ip,
        "public_ip": public_ip
    }


SERVER_INFO = get_server_info()

def setup_logging():
    """
    Initialize logging ONCE per Streamlit session.
    Safe for reruns, EXE, and headless execution.
    """
    logger = logging.getLogger("harmony")
    logger.setLevel(logging.INFO)

    # --- Prevent duplicate handlers on Streamlit reruns ---
    if getattr(logger, "_harmony_initialized", False):
        return logger

    logger.handlers.clear()

    formatter = logging.Formatter(
        "%(asctime)s | %(levelname)s | %(message)s"
    )

    # --- Console logging ---
    try:
        console = logging.StreamHandler()
        console.setFormatter(formatter)
        logger.addHandler(console)
    except Exception:
        pass

    # --- File logging (ROTATING, SAFE) ---
    try:
        log_dir = get_harmony_dir()
        log_dir.mkdir(parents=True, exist_ok=True)

        log_file = log_dir / "harmony.log"

        file_handler = RotatingFileHandler(
            log_file,
            maxBytes=5 * 1024 * 1024,  # 5 MB
            backupCount=3,
            encoding="utf-8"
        )
        file_handler.setFormatter(formatter)
        logger.addHandler(file_handler)

    except Exception:
        # Absolute fallback (TEMP)
        try:
            temp_log = Path(tempfile.gettempdir()) / "harmony.log"
            fallback = RotatingFileHandler(
                temp_log,
                maxBytes=2 * 1024 * 1024,
                backupCount=1,
                encoding="utf-8"
            )
            fallback.setFormatter(formatter)
            logger.addHandler(fallback)
        except Exception:
            pass

    logger.info("Logging initialized successfully")
    logger._harmony_initialized = True
    return logger


# ===============================
# INITIALIZE LOGGER (STREAMLIT SAFE)
# ===============================

if "harmony_logger" not in st.session_state:
    st.session_state.harmony_logger = setup_logging()

logger = st.session_state.harmony_logger


def safe_print(*args):
    """Print + log safely (never crashes, Streamlit / EXE compatible)."""
    msg = " ".join(str(a) for a in args)
    try:
        print(msg)
    except Exception:
        pass
    try:
        logger.info(msg)
    except Exception:
        pass


# ===============================
# VERIFICATION
# ===============================

safe_print("Harmony Engine started successfully")
safe_print("Resolved Harmony dir:", get_harmony_dir())

# === Config: Tesseract ===

import pytesseract
from PIL import Image
import numpy as np
import sys
import io

# ===============================
# TESSERACT CONFIG (Windows + Linux)
# ===============================

import platform
import shutil

if platform.system() == "Windows":

    possible_tesseract = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]

    for p in possible_tesseract:
        if os.path.exists(p):
            pytesseract.pytesseract.tesseract_cmd = p
            break

else:
    # Linux / AWS / Ubuntu
    tesseract_bin = shutil.which("tesseract")

    if not tesseract_bin:
        tesseract_bin = "tesseract"
        print(
            "[WARN] Tesseract binary not found. "
            "OCR will be unavailable until tesseract-ocr is installed."
        )

    pytesseract.pytesseract.tesseract_cmd = tesseract_bin

print(f"[OK] Tesseract detected at: {pytesseract.pytesseract.tesseract_cmd}")


# === Config: Poppler ===
# ===============================
# POPPLER CONFIG
# ===============================

if platform.system() == "Windows":
    POPPLER_PATH = os.path.join(
        BASE_DIR,
        "poppler",
        "poppler-25.07.0",
        "Library",
        "bin"
    )
else:
    POPPLER_PATH = None

# === Config: Font (DejaVuSans.ttf) ===
FONT_PATH = os.path.join(BASE_DIR, "DejaVuSans.ttf")
prop = fm.FontProperties(fname=FONT_PATH)

# === Config: Ollama ===
# (assumes ollama.exe will be installed in the same folder as Harmony.exe)
# === Config: Ollama ===
# ===============================
# OLLAMA CONFIG
# ===============================

if platform.system() == "Windows":

    possible_paths = [
        os.path.join(BASE_DIR, "ollama.exe"),
        r"C:\Program Files\Ollama\ollama.exe",
        os.path.expanduser(
            "~\\AppData\\Local\\Programs\\Ollama\\ollama.exe"
        )
    ]

else:

    possible_paths = [
        "/usr/bin/ollama",
        "/usr/local/bin/ollama",
    ]

OLLAMA_PATH = None

for p in possible_paths:
    if os.path.exists(p):
        OLLAMA_PATH = p
        break

if OLLAMA_PATH:
    print(f"[OK] Ollama detected: {OLLAMA_PATH}")
else:
    print("[INFO] Ollama not installed")	


# === Example usage check ===
print("Tesseract path:", pytesseract.pytesseract.tesseract_cmd)
print("Poppler path:", POPPLER_PATH)
print("Font path:", FONT_PATH)
print("Ollama path:", OLLAMA_PATH)

def safe_write(
    path,
    content
):

    backup = path.with_suffix(".bak")

    try:

        if path.exists():

            shutil.copy2(
                path,
                backup
            )

        temp = path.with_suffix(".tmp")

        temp.write_text(
            content,
            encoding="utf-8"
        )

        temp.replace(path)

        return True

    except Exception as e:

        print(e)

        return False


import hashlib
import builtins
from pathlib import Path

def get_resource_path(relative_path):

    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = BASE_DIR

    return os.path.join(base_path, relative_path)

dejavu_path = get_resource_path("DejaVuSans.ttf")
if os.path.exists(dejavu_path):
    fm.fontManager.addfont(dejavu_path)
    plt.rcParams['font.family'] = 'DejaVu Sans'


DEFAULT_API_KEY = st.secrets[
    "DEFAULT_API_KEY"
]
DEFAULT_FROM = st.secrets[
    "DEFAULT_FROM"
]
CONFIG_FILE = get_harmony_dir() / "email_config.json"

DEFAULT_EMAIL_CONFIG = {
    "mode": "default",
    "recipients": [],
    "api_key": "",
    "from_email": "",
    "smtp_email": "",
    "smtp_password": "",
    "user_email": ""
}

def load_config():
    try:
        if CONFIG_FILE.exists():
            data = json.loads(CONFIG_FILE.read_text(encoding="utf-8"))
            return {**DEFAULT_EMAIL_CONFIG, **data}
    except Exception as e:
        print(f"Email config load error: {e}")

    return DEFAULT_EMAIL_CONFIG.copy()

def save_config(data):
    CONFIG_FILE.parent.mkdir(parents=True, exist_ok=True)
    CONFIG_FILE.write_text(json.dumps(data, indent=2), encoding="utf-8")

config = load_config()

def is_valid_email(email):
    return re.match(r"[^@]+@[^@]+\.[^@]+", email)

import streamlit as st
import sys
from pathlib import Path

# === Page Config ===

st.set_page_config(
    page_title="🧠 Harmony Strategic Assistant",
    layout="wide"
)

# =========================================================
# HARMONY CONSENT SCREEN
# =========================================================

if "consent_accepted" not in st.session_state:
    st.session_state.consent_accepted = False

if not st.session_state.consent_accepted:

    st.markdown("""
    # 🧠 Harmony Strategic Assistant

    Welcome to Harmony AI.

    Please review and accept the Terms & Conditions before continuing.

    🔗 Terms:
    https://harmony-business-ai.xyz/terms
    """)

    st.info("You must accept the Terms & Conditions to continue.")

    if st.checkbox("I accept the Terms & Conditions"):

        st.session_state.consent_accepted = True

        st.rerun()

    st.stop()



# === Global CSS (Enterprise / Big-4 style) ===
st.markdown(
    """
    <style>

    /* ===== Global Layout & Typography ===== */
    html, body, [data-testid="stAppViewContainer"] {
        overflow: auto !important;
        height: auto !important;
        font-family: "DejaVu Sans", "Segoe UI", Roboto, Arial, sans-serif;
        font-size: 15px;
        line-height: 1.55;
        color: #1f2937;
    }

    /* ===== Structured Headings (Enterprise Blue) ===== */
    h1 {
        color: #1f4fd8;
        font-weight: 700;
    }

    h2 {
        color: #1f4fd8;
        font-weight: 600;
    }

    h3 {
        color: #1f4fd8;
        font-weight: 600;
    }

    h4, h5, h6 {
        color: #1f4fd8;
        font-weight: 500;
    }

    /* ===== Heading Spacing ===== */
    h1, h2, h3, h4 {
        margin-top: 1.2em;
        margin-bottom: 0.6em;
    }

    /* ===== Main Content Area ===== */
    .main {
        max-width: 100% !important;
        overflow-x: auto !important;
        padding-left: 2rem;
        padding-right: 2rem;
    }

    /* ===== Sidebar ===== */
    [data-testid="stSidebar"] {
        max-width: 100% !important;
        overflow-x: auto !important;
    }

    /* ===== Tables (Enterprise / Big-4 Style) ===== */
    table {
        border-collapse: collapse;
        font-size: 14px;
    }

    th {
        font-weight: 600;
        background-color: #f9fafb;
    }

    td, th {
        padding: 8px 10px;
        border-bottom: 1px solid #e5e7eb;
    }

    /* ===== Hide Streamlit Default Menu ===== */
    #MainMenu {
        visibility: hidden;
    }

    </style>
    """,
    unsafe_allow_html=True
)


# === HEADER WITH CUSTOM HELP BUTTON ===
col1, col2 = st.columns([6, 1])
with col1:
    st.title("Harmony Strategic Assistant V1.1")
with col2:
    if st.button("❓ Help"):
        st.session_state["show_help"] = True


# === SHOW HELP CONTENT WHEN CLICKED ===
if st.session_state.get("show_help", False):
    st.markdown("---")
    st.markdown("## 📖 Harmony Assistant Help / README")
    st.markdown(read_readme(), unsafe_allow_html=True)
    if st.button("Close Help"):
        st.session_state["show_help"] = False

# === Theme Toggle ===
if "dark_mode" not in st.session_state:
    st.session_state.dark_mode = False

theme = "dark" if st.session_state.dark_mode else "light"

# Apply custom styles only to background + text
st.markdown(f"""
    <style>
    .stApp {{
        background-color: {'#0E1117' if theme == 'dark' else '#FFFFFF'} !important;
        color: {'#FAFAFA' if theme == 'dark' else '#000000'} !important;
    }}
    
    /* Headings */
    h1, h2, h3, h4, h5, h6 {{
        color: {'#FAFAFA' if theme == 'dark' else '#000000'} !important;
    }}

    /* Labels */
    label, .stMarkdown, p, span {{
        color: {'#FAFAFA' if theme == 'dark' else '#000000'} !important;
    }}
    
    /* Input fields */
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea {{
        background-color: {'#1E1E1E' if theme == 'dark' else '#FFFFFF'} !important;
        color: {'#FAFAFA' if theme == 'dark' else '#000000'} !important;
        border: 1px solid {'#555555' if theme == 'dark' else '#CCCCCC'} !important;
        border-radius: 6px;
        padding: 6px;
    }}

    /* Buttons */
    button[kind="primary"] {{
        background-color: {'#444444' if theme == 'dark' else '#F0F0F0'} !important;
        color: {'#FAFAFA' if theme == 'dark' else '#000000'} !important;
        border: 1px solid {'#666666' if theme == 'dark' else '#DDDDDD'} !important;
        border-radius: 8px;
        padding: 6px 12px;
    }}
    </style>
""", unsafe_allow_html=True)

# Toggle button
if st.button("🌗 Toggle Dark/Light Mode"):
    st.session_state.dark_mode = not st.session_state.dark_mode
    

# === Harmony Header ===
st.markdown(
    """
    <div style='background-color:white; padding:12px; border-radius:8px; text-align:center; margin-bottom:15px;'>
        <h2 style='color:#32CD32; margin:0;'>Harmony-GPT</h2>
        <p style='color:#FFFFFF; font-size:14px; margin:0;'>
            💡 All-in-one Management Assistant, designed under patent Arjit's theory for 
            Quantified matrix solution, your refined strategies are just one ping away. <br>
            Under patent theory – powered with RAG & advanced engines <br>
            Designed in India (Vocal for Local) 🇮🇳 <br>
        </p>
    </div>
    """,
    unsafe_allow_html=True
)         
        
    
# =========================================================
# BACKEND DETECTION SYSTEM
# Local + AWS + Production + Fallback
# =========================================================

# =========================================================
# HARMONY STORAGE SYSTEM
# =========================================================

HARMONY_DIR = get_harmony_dir()

# =========================================================
# MAIN DIRECTORIES
# =========================================================


LOG_DIR = HARMONY_DIR / "logs"			
ACTIVITY_LOG = LOG_DIR / "activity.log"
CHAT_LOG_FILE = LOG_DIR / "chat.log"

# =========================================================
# CREATE DIRECTORIES
# =========================================================

for d in [
    HARMONY_DIR,
    LOG_DIR,
]:
    d.mkdir(
        parents=True,
        exist_ok=True
    )



# ============================================	=============
# ACTIVITY LOGGER
# =========================================================

LOG_LOCK = threading.Lock()

def log_user_activity(
    user_id,
    action,
    extra=None
):

    try:

        entry = {

            "timestamp": datetime.datetime.now().isoformat(),

            "user_id": user_id,

            "action": action,

            "server_ip": SERVER_INFO["public_ip"],

            "hostname": SERVER_INFO["hostname"],

            
            "platform": platform.platform(),

            "extra": extra or {}

        }

        with LOG_LOCK:
            with open(
                ACTIVITY_LOG,
                "a",
                encoding="utf-8"
            ) as f:
  
                f.write(
                   json.dumps(entry) + "\n"
                )

    except Exception as e:

        print(f"Activity log error: {e}")

# =========================================================
# BACKEND URLS
# =========================================================

BACKEND_URLS = [

    "http://127.0.0.1:8000",

    os.getenv(
        "BACKEND_URL",
        "https://api.harmony-business-ai.xyz"
    )
]

SYNC_ENDPOINTS = BACKEND_URLS.copy()

# =========================================================
# DEBUG INFO
# =========================================================

print(f"[OK] Harmony Root: {HARMONY_DIR}")


print(f"[OK] Server Mode: {IS_SERVER}")

def is_valid_date(date_str):
    try:
        return datetime.datetime.now().date() <= datetime.datetime.strptime(date_str, "%Y-%m-%d").date()
    except:
        return False



def detect_backend():
    """Detects first available backend from BACKEND_URLS."""
    for url in BACKEND_URLS:
        try:
            if requests.get(f"{url}/ping", timeout=2).status_code == 200:
                return url
        except:
            pass
    return None

BACKEND_URL = st.secrets[
    "BACKEND_URL"
]

# === DAILY SYNC CONTROLS ===
def is_online():
    try:
        socket.create_connection(("8.8.8.8", 53), timeout=2)
        return True
    except:
        return False




import streamlit as st, json, datetime
from pathlib import Path


# === Session State Defaults (AWS / External Auth Bypass) ===

if "user_id" not in st.session_state:
    st.session_state.user_id = USER.get("email", DEFAULT_USER)

if "role" not in st.session_state:
    st.session_state.role = ROLE

active_user = st.session_state.user_id




# ====== Consent Screen ======
consent_accepted = st.checkbox(
    "I acknowledge and accept Harmony usage terms",
    value=True
)


# =========================================================
# HARMONY STREAMLIT CLOUD AUTH SYSTEM
# users.json based
# =========================================================

import json
import bcrypt
import streamlit as st

# =========================================================
# SESSION DEFAULTS
# =========================================================

if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if "username" not in st.session_state:
    st.session_state.username = None

if "role" not in st.session_state:
    st.session_state.role = "viewer"

if "organization" not in st.session_state:
    st.session_state.organization = "Harmony"

import json
import bcrypt
from pathlib import Path

USERS_FILE = Path("users.json")

def create_default_users():

    if USERS_FILE.exists():
        return

    users = {
        "users": [
            {
                "username": "admin",
                "password_hash": bcrypt.hashpw(
                    "Harmony@123".encode(),
                    bcrypt.gensalt()
                ).decode(),
                "role": "admin",
                "organization": "Harmony"
            },

            {
                "username": "user01",
                "password_hash": bcrypt.hashpw(
                    "Europe@123".encode(),
                    bcrypt.gensalt()
                ).decode(),
                "role": "analyst",
                "organization": "Europe"
            },

            {
                "username": "user02",
                "password_hash": bcrypt.hashpw(
                    "China@123".encode(),
                    bcrypt.gensalt()
                ).decode(),
                "role": "analyst",
                "organization": "China"
            }
        ]
    }

    with open(
        USERS_FILE,
        "w",
        encoding="utf-8"
    ) as f:

        json.dump(
            users,
            f,
            indent=4
        )

create_default_users()

# =========================================================
# USERS.JSON LOADER
# =========================================================

USERS_FILE = "users.json"

def load_users():

    try:

        with open(
            USERS_FILE,
            "r",
            encoding="utf-8"
        ) as f:

            return json.load(f)

    except:

        return {"users": []}

# =========================================================
# AUTHENTICATION
# =========================================================

def authenticate_user(
    username,
    password
):

    users_data = load_users()

    for user in users_data["users"]:

        if user["username"] == username:

            try:

                if bcrypt.checkpw(
                    password.encode(),
                    user["password_hash"].encode()
                ):

                    return user

            except:

                pass

    return None

# =========================================================
# LOGIN SCREEN
# =========================================================

if not st.session_state.authenticated:

    st.title("🔐 Harmony Login")

    username = st.text_input(
        "User ID"
    )

    password = st.text_input(
        "Password",
        type="password"
    )

    if st.button("Login"):

        user = authenticate_user(
            username,
            password
        )

        if user:

            st.session_state.authenticated = True

            st.session_state.username = user["username"]

            st.session_state.role = user.get(
                "role",
                "viewer"
            )

            st.session_state.organization = user.get(
                "organization",
                "Harmony"
            )

            st.success(
                f"Welcome {username}"
            )

            st.rerun()

        else:

            st.error(
                "Invalid credentials"
            )

    st.stop()

# =========================================================
# ACTIVE USER
# =========================================================

USER = {

    "email": st.session_state.username,

    "identity": st.session_state.organization

}

ROLE = st.session_state.role

# =========================================================
# ROLE ACCESS
# =========================================================

if ROLE == "admin":

    st.sidebar.success(
        "👑 Admin Access"
    )

elif ROLE == "analyst":

    st.sidebar.info(
        "📊 Analyst Access"
    )

else:

    st.sidebar.warning(
        "👁️ View Only"
    )

# =========================================================
# USER INFO
# =========================================================

st.sidebar.write(
    f"User: {st.session_state.username}"
)

st.sidebar.write(
    f"Organization: {st.session_state.organization}"
)

# =========================================================
# LOGOUT
# =========================================================

if st.sidebar.button(
    "Logout"
):

    st.session_state.clear()

    st.rerun()
 
import smtplib
import requests
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication

def send_resend(to_email, subject, body, attachment=None, filename="Report.pdf", api_key=None, from_email=None):
    """Handles both Brevo and Resend based on the API Key prefix."""
    # Safety: If no API key is provided, we cannot send
    if not api_key:
        return False

    is_brevo = api_key.startswith("xkeysib")
    
    # Prepare Attachment (Base64 is required for most APIs)
    b64_content = ""
    if attachment:
        b64_content = base64.b64encode(attachment).decode('utf-8')

    if is_brevo:
        # --- BREVO API FLOW ---
        payload = {
            "sender": {"name": "Harmony AI", "email": from_email},
            "to": [{"email": to_email}],
            "subject": subject,
            "htmlContent": body
        }
        # Include Reply-To if the user has provided their personal email
        if config.get("user_email"):
            payload["replyTo"] = {"email": config["user_email"]}

        if attachment:
            payload["attachment"] = [{"content": b64_content, "name": filename}]
        
        headers = {"api-key": api_key, "Content-Type": "application/json"}
        res = requests.post("https://api.brevo.com/v3/smtp/email", headers=headers, json=payload)
    else:
        # --- RESEND API FLOW ---
        payload = {
            "from": from_email,
            "to": [to_email],
            "subject": subject,
            "html": body
        }
        if config.get("user_email"):
            payload["reply_to"] = config["user_email"]

        if attachment:
            # Resend takes a slightly different format for attachments
            payload["attachments"] = [{"content": b64_content, "filename": filename}]
            
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        res = requests.post("https://api.resend.com/emails", headers=headers, json=payload)

    return res.status_code in [200, 201]

def send_smtp(to_email, subject, body, attachment=None, filename="Report.pdf"):
    """Handles Standard SMTP (Gmail/Outlook) using UI-entered credentials."""
    smtp_user = config.get("smtp_email")
    smtp_pass = config.get("smtp_password")
    
    if not smtp_user or not smtp_pass:
        return False

    # Auto-detect host
    host = "smtp.gmail.com" if "gmail.com" in smtp_user.lower() else "smtp.office365.com"

    msg = MIMEMultipart()
    msg["Subject"] = subject
    msg["From"] = smtp_user
    msg["To"] = to_email
    msg.attach(MIMEText(body, 'plain'))

    if attachment:
        part = MIMEApplication(attachment, Name=filename)
        part['Content-Disposition'] = f'attachment; filename="{filename}"'
        msg.attach(part)

    try:
        with smtplib.SMTP(host, 587) as s:
            s.starttls()
            s.login(smtp_user, smtp_pass)
            s.send_message(msg)
        return True
    except Exception as e:
        print(f"SMTP Error: {e}")
        return False

def send_email(to_email, subject, body, attachment=None, filename="Report.pdf"):
    """Main wrapper that respects UI settings above all else."""
    mode = config.get("mode")

    if mode == "smtp":
        return send_smtp(to_email, subject, body, attachment=attachment, filename=filename)
    
    elif mode == "custom_api":
        # This block NOW EXCLUSIVELY uses your UI-entered credentials
        return send_resend(
            to_email, subject, body, 
            attachment=attachment, 
            filename=filename,
            api_key=config.get("api_key"), 
            from_email=config.get("from_email")
        )
    else:
        # Default Fallback (System Default)
        return send_resend(
            to_email, subject, body, 
            attachment=attachment, 
            filename=filename,
            api_key=DEFAULT_API_KEY, 
            from_email=DEFAULT_FROM
        )


import pdfplumber
from docx import Document
from pdf2image import convert_from_path
import pandas as pd 

import cv2
import numpy as np
from PIL import Image
import streamlit as st

def ocr_image(img):
    try:
        # 1. [INTACT] Convert PIL Image to a NumPy array
        img_np = np.array(img)

        # 2. [FIXED] Universal channel handling (Prevents crash on grayscale/RGBA)
        if len(img_np.shape) == 2:
            img_cv = img_np
        elif img_np.shape[-1] == 4:
            img_cv = cv2.cvtColor(img_np, cv2.COLOR_RGBA2GRAY)
        else:
            img_cv = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)

        # 3. [INTACT] Apply adaptive thresholding to remove noise
        img_cv = cv2.adaptiveThreshold(
            img_cv, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY, 31, 2
        )

        # 4. [INTACT] Resize for improved OCR accuracy
        img_cv = cv2.resize(img_cv, (1200, 1600), interpolation=cv2.INTER_CUBIC)

        # 5. [IMPROVED] Use image_to_data for 90%+ Table Accuracy
        # This keeps columns aligned. We still use your custom_config.
        custom_config = r'--oem 3 --psm 6'
        data = pytesseract.image_to_data(img_cv, config=custom_config, output_type=pytesseract.Output.DATAFRAME)

        # 6. [NEW] Filter out low-confidence noise (Cleans up the "|" and "i" artifacts)
        df = data[data.conf > 10].dropna()
        
        # 7. [NEW] Reconstruct lines to maintain table structure
        if not df.empty:
            lines = df.groupby('line_num')['text'].apply(lambda x: ' '.join(x)).tolist()
            text = "\n".join(lines)
        else:
            # Fallback to your original string method if grouping fails
            processed_img = Image.fromarray(img_cv)
            text = pytesseract.image_to_string(processed_img, config=custom_config)

        # 8. [INTACT] Clean up extra spaces
        text = " ".join(text.split())

        return text

    except Exception as e:
        st.error(f"❌ Image OCR failed: {e}")
        return ""



# --- Helper: OCR for PDFs ---
def ocr_pdf(pdf_file):
    full_text = ""
    try:
        # First try pdfplumber for structured text & tables
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    full_text += page_text + "\n\n"

        # If no text found → likely scanned → OCR each page
        if not full_text.strip():
            if POPPLER_PATH:
                pages = convert_from_path(
                    pdf_file,
                    dpi=150,
                    poppler_path=POPPLER_PATH
                )  
            else:
               pages = convert_from_path(
                    pdf_file,
                    dpi=150
                )
            for i, page in enumerate(pages):
                page = page.resize((1200, 1600))
                text = ocr_image(page)
                full_text += f"\n--- Page {i+1} ---\n{text}"

    except Exception as e:
        st.warning(f" PDF OCR failed: {e}")
    return full_text

# === Streamlit Tabs ===
tab1, tab2 = st.tabs(["📄 File Upload", "🖊️ Manual Input"])

with tab1:
    uploaded_file = st.file_uploader(
        "Upload a file",
        type=["pdf", "docx", "txt", "jpg", "jpeg", "png", "xlsx"]
    )
    extracted_text = ""

    if uploaded_file:
        file_type = uploaded_file.name.split(".")[-1].lower()

        if file_type == "pdf":
            extracted_text = ocr_pdf(uploaded_file)

        elif file_type == "docx":
            doc = Document(uploaded_file)
            extracted_text = "\n".join([para.text for para in doc.paragraphs])

        elif file_type == "txt":
            extracted_text = uploaded_file.read().decode("utf-8", errors="ignore")

        elif file_type in ["jpg", "jpeg", "png"]:
            img = Image.open(uploaded_file)
            extracted_text = ocr_image(img)

         # ----------------- XLSX (NEW) -----------
        elif file_type == "xlsx":
            try:
                # Read ALL sheets
                xls = pd.ExcelFile(uploaded_file)
                excel_text = ""

                for sheet in xls.sheet_names:
                    df = pd.read_excel(uploaded_file, sheet_name=sheet, dtype=str)
                    df = df.fillna("")  # avoid NaN

                    excel_text += f"\n=== Sheet: {sheet} ===\n"
                    excel_text += df.to_string(index=False)
                    excel_text += "\n\n"

                extracted_text = excel_text.strip()

            except Exception as e:
                st.error(f"❌ Excel Read Error: {e}")
                extracted_text = ""    

        # Automatically populate the Harmony input box
        st.session_state['input_text'] = extracted_text

with tab2:
    manual_input = st.text_area("📝 Enter Text Manually", height=150, key='input_text')

# === Harmony Input Text Area ===
st.subheader("📥 Harmony Input Text")
st.text_area("Input ready for analysis", st.session_state.get('input_text', ''), height=300)
    
# === Comments and Priority ===
col1, col2 = st.columns(2)
with col1:
    user_comment = st.text_input("User Comments")
    user_priority = st.slider("User Priority (%)", 0, 100, 50)

with col2:
    doer_comment = st.text_input("Doer Comments")
    doer_priority = 100 - user_priority


# === Step 1: Select Problem Solving Methodology ===
model_type = st.sidebar.selectbox("📘 Select Problem-Solving Methodology", [
    "5 Whys", "Fishbone", "SWOT", "DMAIC", "BCG", "Porter", "Iceberg", "Pareto", "User-defined" 
])

from dotenv import load_dotenv
import os


# =========================================================

# STREAMLIT CLOUD ENGINE SELECTION

# =========================================================

engine_type = st.sidebar.selectbox(
"🧠 Choose AI Engine",
[
"Flashmind RAG",
"Omnicore RAG",
"Flashmind Market",
"OpenAi-GPT",
"Product-R&D-Omni",
"Product-Optimizations-Omni",
"Omnicore_Finance"
]
)

st.session_state.selected_engine = engine_type

# =========================================================

# STREAMLIT SECRETS

# =========================================================

FLASHMIND_KEY = st.secrets.get(
"FLASHMIND_KEY",
""
)

OMNICORE_KEY = st.secrets.get(
"OMNICORE_KEY",
""
)

OPENAI_API_KEY = st.secrets.get(
"OPENAI_API_KEY",
""
)

BACKEND_URL = st.secrets.get(
"BACKEND_URL",
""
)

# =========================================================

# DEBUG PANEL

# =========================================================

with st.sidebar:

    st.markdown("---")

    st.subheader("🔧 System Status")

    st.write(
        "Selected Engine:",
        st.session_state.selected_engine
    )

    st.write(
        "FLASHMIND_KEY:",
        bool(FLASHMIND_KEY)
    )

    st.write(
        "OMNICORE_KEY:",
        bool(OMNICORE_KEY)
    )

    st.write(
        "OPENAI_API_KEY:",
        bool(OPENAI_API_KEY)
    )

    st.write(
        "BACKEND_URL:",
        bool(BACKEND_URL)
    )

# =========================================================
# API KEY RESOLUTION
# =========================================================

GROQ_API_KEY = FLASHMIND_KEY

OPENROUTER_API_KEY = OMNICORE_KEY

OPENAI_KEY = OPENAI_API_KEY

# =========================================================
# ENGINE ROUTING
# =========================================================

if engine_type in [
    "Flashmind RAG",
    "Flashmind Market"
]:

    ACTIVE_API_KEY = GROQ_API_KEY

elif engine_type in [
    "Omnicore RAG",
    "Product-R&D-Omni",
    "Product-Optimizations-Omni",
    "Omnicore_Finance"
]:

    ACTIVE_API_KEY = OPENROUTER_API_KEY

elif engine_type == "OpenAi-GPT":

    ACTIVE_API_KEY = OPENAI_KEY

else:

    ACTIVE_API_KEY = ""

# =========================================================
# STATUS DISPLAY
# =========================================================

if engine_type in [
    "Flashmind RAG",
    "Flashmind Market"
]:

    if GROQ_API_KEY:

        st.sidebar.success(
            "✅ Flashmind Connected"
        )

    else:

        st.sidebar.error(
            "❌ FLASHMIND_KEY Missing"
        )

elif engine_type in [
    "Omnicore RAG",
    "Product-R&D-Omni",
    "Product-Optimizations-Omni",
    "Omnicore_Finance"
]:

    if OPENROUTER_API_KEY:

        st.sidebar.success(
            "✅ Omnicore Connected"
        )

    else:

        st.sidebar.error(
            "❌ OMNICORE_KEY Missing"
        )

elif engine_type == "OpenAi-GPT":

    if OPENAI_KEY:

        st.sidebar.success(
            "✅ OpenAI Connected"
        )

    else:

        st.sidebar.error(
            "❌ OPENAI_API_KEY Missing"
        )

# =========================================================
# FINAL SESSION VALUES
# =========================================================

st.session_state["engine"] = engine_type

st.session_state["active_api_key"] = ACTIVE_API_KEY

st.session_state["backend_url"] = BACKEND_URL

st.sidebar.header("📧 Email Settings")

mode = st.sidebar.radio(
    "Mode",
    ["default", "custom_api", "smtp"],
    index=["default", "custom_api", "smtp"].index(config.get("mode", "default"))
)

user_email = st.sidebar.text_input("Reply-To Email", value=config.get("user_email", ""))

api_key = config.get("api_key", "")
from_email = config.get("from_email", "")

if mode == "custom_api":
    api_key = st.sidebar.text_input("API Key", type="password", value=api_key)
    from_email = st.sidebar.text_input("From Email", value=from_email)

smtp_email = config.get("smtp_email", "")
smtp_password = config.get("smtp_password", "")

if mode == "smtp":
    smtp_email = st.sidebar.text_input("SMTP Email", value=smtp_email)
    smtp_password = st.sidebar.text_input("App Password", type="password", value=smtp_password)

recipients = st.sidebar.text_area("Recipients (comma separated)", value=",".join(config.get("recipients", [])))

if st.sidebar.button("Save Email Settings"):
    config.update({
        "mode": mode,
        "api_key": api_key,
        "from_email": from_email,
        "smtp_email": smtp_email,
        "smtp_password": smtp_password,
        "user_email": user_email,
        "recipients": [e.strip() for e in recipients.split(",") if e.strip()]
    })
    save_config(config)
    st.sidebar.success("Saved")

import subprocess, re, shutil, streamlit as st

possible_paths = [
    os.path.join(BASE_DIR, "ollama.exe"),
    r"C:\Program Files\Ollama\ollama.exe",
    os.path.expanduser("~\\AppData\\Local\\Programs\\Ollama\\ollama.exe")
]

      
import datetime
from typing import List


# === API Config & Daily Limits ===
GOOGLE_API_KEYS = st.secrets(
    "GOOGLE_API_KEYS"
).split(",")

GOOGLE_CSE_ID = "236166c7fd4994632"
GOOGLE_DAILY_LIMIT = 30

BING_API_KEY = st.secrets[
    "BING_API_KEY"
]
BING_API_HOST = "bing-web-search1.p.rapidapi.com"
BING_DAILY_LIMIT = 5

YAHOO_API_KEY = st.secrets[
    "YAHOO_API_KEY"
]
YAHOO_API_HOST = "apidojo-yahoo-finance-v1.p.rapidapi.com"
YAHOO_DAILY_LIMIT = 5

# === Offline References ===
OFFLINE_REFERENCES = [
    "https://www.forbes.com/",
    "https://www.bbc.com/news/business",
    "https://www.economist.com/business/",
    "https://www.investopedia.com/"
]

# --- Helpers ---
def _reset_daily_counter(key: str):
    today = datetime.date.today().isoformat()
    if key not in st.session_state or st.session_state[key]["date"] != today:
        st.session_state[key] = {"date": today, "count": 0}

# --- Search Engines ---
def fetch_google(query: str) -> List[str]:
    _reset_daily_counter("google_calls")
    if st.session_state["google_calls"]["count"] >= GOOGLE_DAILY_LIMIT:
        return []
    for key in GOOGLE_API_KEYS:  # try multiple keys
        try:
            url = "https://www.googleapis.com/customsearch/v1"
            params = {"key": key, "cx": GOOGLE_CSE_ID, "q": query, "num": 3}
            resp = requests.get(url, params=params, timeout=10)
            resp.raise_for_status()
            items = resp.json().get("items", [])
            links = [item.get("link") for item in items if item.get("link")]
            if links:
                st.session_state["google_calls"]["count"] += 1
                return links
        except Exception:
            continue
    return []
 
def fetch_bing(query: str) -> List[str]:
    _reset_daily_counter("bing_calls")
    if st.session_state["bing_calls"]["count"] >= BING_DAILY_LIMIT:
        return []
    try:
        url = f"https://{BING_API_HOST}/search"
        headers = {"X-RapidAPI-Key": BING_API_KEY, "X-RapidAPI-Host": BING_API_HOST}
        params = {"q": query, "count": "3"}
        resp = requests.get(url, headers=headers, params=params, timeout=10)
        resp.raise_for_status()
        data = resp.json()
        links = [item["url"] for item in data.get("webPages", {}).get("value", [])]
        if links:
            st.session_state["bing_calls"]["count"] += 1
        return links
    except Exception:
        return []

def fetch_yahoo(query: str) -> List[str]:
    if not YAHOO_API_KEY:
        return []
    try:
        url = "https://yahoo-search5.p.rapidapi.com/web/search"
        headers = {"X-RapidAPI-Key": YAHOO_API_KEY, "X-RapidAPI-Host": "apidojo-yahoo-finance-v1.p.rapidapi.com"}
        params = {"q": query, "count": "3"}
        resp = requests.get(url, headers=headers, params=params, timeout=10)
        resp.raise_for_status()
        data = resp.json()
        return [item["link"] for item in data.get("web", {}).get("results", [])]
    except Exception:
        return []

# --- Main Orchestrator ---
def get_references(query: str, online_mode: bool = True) -> List[str]:
    """
    Get reference links for the query.
    Tries Google → Bing → Yahoo if online_mode=True.
    Falls back to offline references otherwise.
    """
    if online_mode:
        for fetcher in (fetch_google, fetch_bing, fetch_yahoo):
            links = fetcher(query)
            if links:
                return links

    # Show "trying references" only once
    if "references_info_shown" not in st.session_state:
        st.info("trying references")
        st.session_state["references_info_shown"] = True

    return OFFLINE_REFERENCES

# === Build Final Prompt ===
def build_f_prompt(user_comment, doer_comment, final_input, model_type, user_priority, perspective="user"):
    """
    Construct a comprehensive financial analysis prompt.
    Ensures multi-year tables, ratio analysis, vertical/horizontal/variance/valuation analysis,
    qualitative/contextual insights, forward-looking guidance, suggestions, and visual aids.
    
    Works for final input, user comments, or doer comments.
    """
    
    # --- Combine all inputs into context ---
    combined_input = ""
    if final_input:
        combined_input += f"Business Context: {final_input}\n\n"
    if user_comment:
        combined_input += f"User Input: {user_comment}\n\n"
    if doer_comment:
        combined_input += f"Doer Input: {doer_comment}\n\n"
    
    combined_input += f"Model Type: {model_type}\nUser Priority: {user_priority}\n"

    # --- Role/system instruction ---
    perspective = str(perspective).lower()
    
    if perspective == "doer":
        role_instruction = (
            "You are a professional financial analyst. Prepare a detailed financial report current 2026 from a Doer perspective. "
            "The report must include the following sections:\n\n"
            "1. Overall Financial Health – liquidity, profitability, solvency, operational efficiency.\n"
            "2. Ratio Analysis – key ratios (ROE, ROA, net margin, operating margin, current ratio, quick ratio, debt-to-equity, interest coverage, EPS, cash-on-hand, CapEx) benchmarked against historical data and industry standards.\n"
            "3. Vertical Analysis (Common-Size) – each line item as a % of total revenue or assets.\n"
            "4. Horizontal Analysis (Trend) – multi-period comparisons showing % changes and growth patterns.\n"
            "5. Variance Analysis (Budget vs Actual) – highlight differences and explain causes.\n"
            "6. Valuation Analysis – intrinsic value (DCF, comparable company analysis) with investment insights.\n"
            "7. Qualitative & Contextual Insights – MD&A, industry/economic factors, risk assessment, notes/disclosures.\n"
            "8. Forward-Looking Guidance – forecasts, projections, and actionable recommendations for operational improvements.\n"
            "9. Suggestions & Recommendations (latest-2026) for Major Changes – highlight critical areas needing attention.\n"
            "10. Visual Aids & Data Visualization – propose tables, charts, and dashboards for clarity.\n\n"
            "Use headings, subheadings, bullet points, numeric examples, and tables. "
            "Ensure clarity, depth, and actionable insights.\n\n"
            "OUTPUT RULES:\n"
            "- Use structured headings and numbered sections\n"
            "- Keep paragraphs concise (maximum 3 lines each)\n"
            "- Use tables where comparisons add clarity\n"
            "- Maintain a neutral, professional engineering tone\n"
            "- No emojis, no casual language, no marketing hype\n"
            "- Use precise, audit-safe terminology suitable for technical and leadership review."
        )
    else:
        role_instruction = (
            "You are a professional financial analyst. Prepare a detailed financial report current 2026 from a User perspective. "
            "The report must include the following sections:\n\n"
            "1. Overall Financial Health – liquidity, profitability, solvency, operational efficiency.\n"
            "2. Ratio Analysis – key ratios benchmarked against historical data and industry norms.\n"
            "3. Vertical Analysis (Common-Size) – each line item as a % of total revenue or assets.\n"
            "4. Horizontal Analysis (Trend) – multi-period comparisons showing % changes and growth patterns.\n"
            "5. Variance Analysis (Budget vs Actual) – highlight differences and explain causes.\n"
            "6. Valuation Analysis – intrinsic value (DCF, comparable company analysis) with strategic insights.\n"
            "7. Qualitative & Contextual Insights – MD&A, industry/economic factors, risk assessment, notes/disclosures.\n"
            "8. Forward-Looking Guidance – projections and actionable recommendations for strategic decision-making.\n"
            "9. Suggestions & Recommendations for Major Changes (latest-2026) – highlight critical areas needing attention.\n"
            "10. Visual Aids & Data Visualization – propose tables, charts, and dashboards to enhance comprehension.\n\n"
            "Use headings, subheadings, bullet points, numeric examples, and tables. "
            "Ensure clarity, depth, and actionable insights.\n\n"
            "OUTPUT RULES:\n"
            "- Use structured headings and numbered sections\n"
            "- Keep paragraphs concise (maximum 3 lines each)\n"
            "- Use tables where comparisons add clarity\n"
            "- Maintain a neutral, professional engineering tone\n"
            "- No emojis, no casual language, no marketing hype\n"
            "- Use precise, audit-safe terminology suitable for technical and leadership review."
        )

    # --- Merge role instruction with combined input ---
    f_prompt = role_instruction + "\n\n" + combined_input
    return f_prompt



# === Prompt Builder (with integrated references + image hooks) ===
def build_prompt(user, doer, input_text, model, prio, online_mode: bool = True):
    context = f"""User Priority: {prio}%\nDoer Priority: {100 - prio}%\n\n📥 User Input:\n{user}\n\n🛠️ Doer Input:\n{doer}\n\n📄 Main Content:\n{input_text}"""

    refs = get_references(input_text, online_mode)
    refs_md = "\n".join([f"- [{url}]({url})" for url in refs])

    definition_section = f"""
Definition of **{input_text}** with authoritative insights (focus on 2026 strategies):

References (auto-expand with text/images if available):  
{refs_md}
"""

    return f"""
Use the *{model}* methodology to analyze the context below.

ok Provide:

Perform a **CEO-level Root Cause Analysis (RCA)** for the given problem,
using **engineering science, chemistry, physics, and real-world industry evidence**.
The analysis must be suitable for **Board review, regulatory audits, and strategic decision-making**.

The response must demonstrate **deep technical reasoning**, not surface-level explanations.


SCOPE & DEPTH REQUIREMENTS (MANDATORY)

- Explain root causes using:
  - Chemistry (chemical reactions, fault mechanisms, physics & material science)
  - Physics (pressure, stress, thermodynamics, fatigue, diffusion)
  - Engineering principles (design, manufacturing, validation, lifecycle)
- Quantify impacts wherever possible
- Link technical failures to **business, safety, and financial consequences**
- Use **2026–2026 industry practices only**


DELIVERABLE STRUCTURE (STRICT)

1. ROOT CAUSE IDENTIFICATION (QUANTIFIED)

- Identify the **primary root cause(s)** and **contributing causes**
- Assign **relevance percentages** to each cause
- Percentages must **sum exactly to 100%**
- Causes must be explained using:
  - Scientific mechanisms (chemistry + physics)
  - Engineering conditions (pressure, materials, environment)

2. DETAILED RECOMMENDATIONS (PARAGRAPH FORMAT)

- Provide **one paragraph per root cause**
- Explain:
  - Why the recommendation works scientifically
  - How it mitigates the failure mechanism
  - Engineering and operational feasibility
  - Avoid repeating table text

3. RCA SUMMARY TABLE (FOR CHARTING)

Provide a markdown table ONLY (no ASCII boxes):

| Root Cause | Contribution (%) | Recommended Solution |
|------------|------------------|----------------------|
| Cause 1    | XX               | Solution summary     |
| Cause 2    | XX               | Solution summary     |
| Cause 3    | XX               | Solution summary     |

4. CHART HEADINGS (REQUIRED)

Include explicit headings for visualization tools:

- Bar Chart: Root Causes Contribution (%)
- Pie Chart: Root Cause Distribution

5. NUMERIC / COMPARATIVE TABLES

- Create **separate tables** for:
  - Percentage contributions
  - Material comparisons
  - Technology alternatives
- These tables must be suitable for Pie or Bar charts

6. DETAILED ENGINEERING & OPERATIONAL SUGGESTIONS

- Design changes
- Material upgrades
- Manufacturing process controls
- Testing & validation improvements
- Service and maintenance interventions

7. IMPLEMENTABLE INDUSTRY EXAMPLES (2026–2026)

- Cite **Top 5 customers / OEMs / industries** that faced similar issues
- Explain:
  - What failure they experienced
  - The root cause identified
  - How they rectified it (engineering + process + policy)
- Examples must be realistic and industry-aligned
  (e.g., automotive OEMs, energy companies, infrastructure operators)

8. AUTHORITATIVE INSIGHTS (2026 CONTEXT)

- Reference **authoritative insights (2026 only)** such as:
  - Industry standards
  - Regulatory shifts
  - OEM R&D directions
  - Testing methodology evolution
- Contextualize references (do not just quote them)
- If references include images:
  - Preserve them inline using markdown:
    ![alt](url) or <img>


CEO-LEVEL EXPECTATIONS

- Clearly explain:
  - Risk exposure
  - Safety implications
  - Financial and reputational impact
- Translate technical failures into **business consequences**
- Highlight what happens **if no action is taken**


IMPORTANT RULES (NON-NEGOTIABLE)

Use a clean enterprise report style similar to Big-4 consulting decks.

Writing & formatting rules:
- Inter-style professional typography (clean, modern, readable – UI controlled)
- Structured headings with numbering
- Short paragraphs (maximum 3 lines)
- Tables for comparisons where relevant
- Bold key figures and conclusions
- Analyst / CFO briefing tone
- Explain implications, not just raw data
- No emojis, no casual language
- High clarity, audit-safe wording
- Assume outputs may be used in Board decks, audits, or regulatory reviews

{definition_section}

⚠ Ensure:  
- Use percentages or numeric values for contributions (required for charts).  
- Use *markdown tables* (pipes |) rather than ASCII boxes.  
- Include headings for Pie or Bar charts so your renderer can pick them up.  

📌 Context:  
{context}  

(We understand the complexity of problems and harmony required for solution-oriented decisions, Arjit's Theory of Problem Solving under patent: with IPI India)  
"""

# === Streamlit Runner ===
def run_analysis_model(user_query, user, doer, model, prio, online_mode=True):
    # Build the unified prompt
    prompt = build_prompt(user, doer, user_query, model, prio, online_mode)

    # Show references separately in UI (clickable)
    st.subheader(" References")
    refs = get_references(user_query, online_mode)
    for ref in refs:
        st.markdown(f"- [Source]({ref})")

        # Auto-display if reference looks like an image link
        if ref.lower().endswith((".png", ".jpg", ".jpeg", ".gif", ".webp")):
            st.image(ref, caption="Inline article image")

    return prompt

# === Build Final General Market Prompt with References ===
def build_market_prompt(user_comment, doer_comment, final_input, market_focus="general", perspective="user", online_mode=True):
    """
    Construct the final general market analysis prompt (m_prompt) for Flashmind Market, current 2026, under 530 words.
    Focuses on market insights: opportunities, risks, customer needs, trends, competition, and strategies.
    Combines user & doer perspectives with a structured market analysis outline.
    Integrates references (text/images) dynamically using get_references.
    """
    
    # --- Enterprise style rules ---
    style_rules = """
IMPORTANT RULES – OUTPUT STYLE & QUALITY

Use a clean enterprise report style equivalent to Big-4 consulting decks.

Writing & formatting requirements:
- Professional, modern, enterprise-readable typography (UI-controlled)
- Clearly numbered section headings
- Short paragraphs (maximum 3 lines each)
- Tables for comparisons, benchmarks, and summaries
- Bold only key figures and conclusions
- Analyst / CFO briefing tone
- Explain implications, not just raw data
- No emojis or casual language
- High clarity with audit-safe wording
- Assume output may be used in Board decks, audits, or regulatory reviews
""".strip()
    
    # --- Raw input layer ---
    q_prompt = (
        f"User Input: {user_comment}\n\n"
        f"Doer Input: {doer_comment}\n\n"
        f"Market Context: {final_input}\n\n"
        f"Market Focus: {market_focus}\n"
    )

    # --- Fetch references dynamically ---
    refs = get_references(final_input, online_mode)
    refs_md = "\n".join([f"- [{url}]({url})" for url in refs])
    # Also preserve images inline
    image_refs_md = "\n".join([f"![Image]({url})" for url in refs if url.lower().endswith((".png", ".jpg", ".jpeg", ".gif", ".webp"))])

    # --- Role/system instruction ---
    perspective = str(perspective).lower()
    if perspective == "doer":
        role_instruction = (
            "You are a market strategist. Analyze the scenario from a Doer perspective, under 500 words. "
            "Provide practical, step-by-step actions: go-to-market plan, customer acquisition, "
            "pricing, operations scaling, and execution strategies. "
            "Include market tables, charts, and competitor benchmarks."
        )
    else:
        role_instruction = (
            "You are a market strategist. Analyze the scenario from a User perspective, under 500 words. "
            "Focus on market trends, opportunities, customer demand, adoption barriers, "
            "and long-term sustainability. Use structured insights, clear headings, "
            "and include summary market tables or charts."
        )

    # --- Structured Market Analysis Outline ---
    structure = """
Format the report strictly as follows:

I. Executive Summary
  A. Overview of the market analysis report 2026 
  B. Key findings and recommendations 2026
II. Introduction
  A. Purpose and objectives of the report
  B. Background information on the market being analyzed
  C. Scope of the analysis
III. Methodology
  A. Description of the research methods used for data collection
  B. Sources of data (primary and secondary)
  C. Limitations of the analysis
IV. Market Overview
  A. Definition and classification of the market
  B. Market size with actual month and year with references and growth rate
  C. Market trends with top 5 customers of products with turnover sales and intel in detail and drivers
  D. Market challenges and barriers
  E. Market opportunities
V. Market Segmentation
  A. Description of the different market segments
  B. Analysis of each segment's size, growth, and potential
  C. Identification of key customer groups within each segment
VI. Competitive Analysis
  A. Identification and profiling of key competitors
  B. Assessment of their market share and positioning
  C. Analysis of their strengths, weaknesses, opportunities, and threats (SWOT)
  D. Comparison of competitor products, pricing, distribution channels and market strategies
VII. Customer Analysis
  A. Identification of target customers and their characteristics
  B. Assessment of customer needs, preferences, and behaviors
  C. Analysis of customer buying patterns and decision-making process
VIII. Market Entry Strategy & Sales
  A. Evaluation of potential market entry barriers and risks
  B. Assessment of market entry modes (e.g., exporting, licensing, joint ventures)
  C. Recommendation of the most suitable market entry strategy with existing and recommended marketing strategies 
  D. Sales Leads, Ideas, Client generation, links for sure shot sales, types of sales recommended top 5 industries and 6 -7 clents each of industry potential clients  
IX. Market Forecast & Sales
  A. Projection of Actual current market size with references and growth in the short and long term
  B. Analysis of factors influencing the market forecast
  C. Estimation of market demand and sales potential
  D. How to get more clients and easy channel to approach and top 10 customers with consumptions to approach 
  E. reports and leads and links for sales which can be fetched to explore sales business opportunity and reach to make customers with top 5 industries and 6 -7 clents each of industry potential clients
X. Conclusion
  A. Summary of the key findings from the market analysis
  B. Recommendations for market entry, Pre and post Brand Strategy 360 degree, Market Communications recommendations, and future actions with company specific examples
XI. Appendices
  A. Supporting data and charts
  B. Glossary of terms
  C. References and sources used in the analysis Current Year 2026 
"""

    # --- Definition/References Section ---
    definition_section = f"""
Definition & References for Market Context:

{refs_md}

Inline Images (if any):

{image_refs_md}
"""

    # --- Merge all layers into final m_prompt ---
    m_prompt = role_instruction + "\n\n" + q_prompt + "\n\n" + structure + "\n\n" + definition_section
    return m_prompt

import re
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import streamlit as st

# === Clean numeric columns safely ===
def _clean_numeric_columns(df):
    def extract_numeric(value):
        try:
            if isinstance(value, (int, float)):
                return float(value)
            if isinstance(value, str):
                clean_val = (
                    value.strip()
                    .replace(",", "")
                    .replace("%", "")
                    .replace("percent", "")
                    .replace("Percent", "")
                )
                match = re.search(r"[-+]?\d*\.?\d+", clean_val)
                if match:
                    return float(match.group())
            return np.nan
        except:
            return np.nan

    for col in df.columns:
        try:
            df.loc[:, col] = df[col].apply(extract_numeric)
            if df[col].notna().mean() < 0.5:
                df.loc[:, col] = df[col].astype(str)
        except:
            df[col] = df[col].astype(str)
    return df

# === Parse Markdown table, inline percentages, numbered lists ===
def _parse_markdown_table(text: str):
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    data = []

    # 1️⃣ Markdown table
    if any('|' in line for line in lines):
        try:
            table_lines = [line for line in lines if not re.match(r'^\|[-:\s|]+\|$', line)]
            if len(table_lines) >= 2:
                rows = [re.split(r'\s*\|\s*', l.strip('|')) for l in table_lines]
                df = pd.DataFrame(rows[1:], columns=rows[0])
                df.iloc[:, 0] = df.iloc[:, 0].astype(str)
                if df.shape[1] > 1:
                    df.iloc[:, 1:] = _clean_numeric_columns(df.iloc[:, 1:])
                return df
        except:
            pass

    # 2️⃣ Inline percentages / numbered lists
    for line in lines:
        match = re.match(r"(?:\d+\.\s*)?[\*\-]?\s*([\w\s&]+)\s*[:\(]\s*(\d+)%", line)
        if match:
            cause, percent = match.groups()
            data.append([cause.strip(), float(percent)])
    if data:
        return pd.DataFrame(data, columns=["Category", "Value"])

    # 3️⃣ Fallback: any line with number
    for line in lines:
        numbers = re.findall(r"[-+]?\d*\.?\d+", line)
        if numbers:
            text_part = re.sub(r"[-+]?\d*\.?\d+%?", "", line).strip()
            if text_part:
                data.append([text_part] + [float(n) for n in numbers])
    if data:
        df = pd.DataFrame(data)
        if df.shape[1] == 2:
            df.columns = ["Category", "Value"]
        return df

    return None

# === Shorten labels (max 5 words) ===
def _short_labels(labels, max_words=5):
    return [" ".join(str(l).split()[:max_words]) for l in labels]

# === Bar chart helper ===
def _plot_bar_chart(df, cat_col, num_cols, title, model_name):
    if len(df) <= 1:
        return None  # 🚫 Skip single-bar charts

    # 🚫 Remove rows where all numeric values are 0 or NaN
    df = df.dropna(subset=num_cols, how="all")
    df = df[(df[num_cols].sum(axis=1) != 0)]

    # 🚫 Remove categories that look like URLs or junk
    df = df[~df[cat_col].str.contains(r"(http|ftp|www|\.pdf)", case=False, na=False)]
    df = df[df[cat_col].str.strip().str.lower().ne("0")]
    df = df[df[cat_col].str.strip().str.lower().ne("nan")]

    # [ok] Keep only top 10 categories
    if len(df) > 10:
        df = df.sort_values(num_cols[0], ascending=False).head(10)

    if df.empty:
        return None

    # 🔹 Fix NaN → 0
    df[num_cols] = df[num_cols].fillna(0)

    # 🔹 Shorten category labels
    df[cat_col] = _short_labels(df[cat_col])

    fig, ax = plt.subplots(figsize=(max(6, len(df[cat_col]) * 0.7), 4))
    x = range(len(df[cat_col]))
    width = 0.25

    for i, col in enumerate(num_cols):
        bars = ax.bar([p + i * width for p in x], df[col], width=width, label=col)

        # [ok] Add numbers above bars
        for bar in bars:
            yval = bar.get_height()
            ax.text(
                bar.get_x() + bar.get_width() / 2,
                yval + 0.5,
                f"{int(yval)}",
                ha="center",
                va="bottom",
                fontsize=9,
                color="black",
            )

    ax.set_xticks([p + width * (len(num_cols) / 2) for p in x])
    ax.set_xticklabels(df[cat_col], rotation=45, ha="right", fontsize=9) # 🔄 Slanted labels
    ax.set_ylabel("Value")
    ax.set_xlabel(cat_col)

    # [ok] Title aligned top-left
    ax.set_title(f"{model_name} · {title}", loc="left")

    ax.legend()
    fig.subplots_adjust(bottom=0.25)
    return fig

# === Render AI output charts automatically ===
def render_charts_from_ai_output(ai_output):
    if isinstance(ai_output, (tuple, list)):
        ai_output = " ".join(map(str, ai_output))

    model_name = st.session_state.get("model", "AI Model")
    rendered_any = False

    # --- JSON chart blocks ---
    json_blocks = re.findall(r"\{.*?\}", ai_output, re.DOTALL)
    for jb in json_blocks:
        try:
            data = json.loads(jb)
            chart = data.get("chart")
            if chart:
                title = chart.get("title", "AI Chart")
                if chart["type"] == "pie":
                    fig, ax = plt.subplots(figsize=(6,6))
                    values = chart.get("data", [])
                    labels = chart.get("labels", [])
                    if len(values) > 1 and sum(values) > 0:
                        values = [v / sum(values) * 100 for v in values]
                        labels = _short_labels(labels)
                        ax.pie(values, labels=labels, autopct='%1.1f%%')
                        ax.set_title(f"{model_name} · {title}", loc="left")
                        plt.tight_layout()
                        st.pyplot(fig)
                elif chart["type"] == "bar":
                    df = pd.DataFrame({"Category": chart.get("x-axis", []),
                                       "Value": chart.get("y-axis", [])})
                    fig = _plot_bar_chart(df, "Category", ["Value"], title, model_name)
                    if fig:
                        st.pyplot(fig)
                rendered_any = True
        except:
            continue

    # --- Tables and lists ---
    parts = [p.strip() for p in re.split(r"\n\s*\n", ai_output) if p.strip()]
    for idx, part in enumerate(parts, start=1):
        df = _parse_markdown_table(part)
        if df is None or df.empty:
            continue

        # Identify category column
        cat_col = next((c for c in df.columns if not pd.api.types.is_numeric_dtype(df[c])), df.columns[0])
        df[cat_col] = df[cat_col].astype(str).ffill()

        num_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        if not num_cols:
            continue

        # Normalize numeric columns to 100
        for col in num_cols:
            total = df[col].sum()
            if total > 0:
                df[col] = df[col] / total * 100

        try:
            # --- Pie chart logic ---
            if 2 <= len(df) <= 5 and 1 <= len(num_cols) <= 2:
                fig, ax = plt.subplots(figsize=(6,6))
                if len(num_cols) == 1:
                    ax.pie(df[num_cols[0]], labels=_short_labels(df[cat_col]), autopct='%1.1f%%')
                else:
                    fig, axes = plt.subplots(1, len(num_cols), figsize=(6*len(num_cols),6))
                    if len(num_cols) == 2:
                        axes = [axes[0], axes[1]]
                    for i, col in enumerate(num_cols):
                        axes[i].pie(df[col], labels=_short_labels(df[cat_col]), autopct='%1.1f%%')
                        axes[i].set_title(col)
                plt.tight_layout()
                st.pyplot(fig)

            # --- Bar chart logic ---
            else:
                fig = _plot_bar_chart(df, cat_col, num_cols, f"Chart {idx}: {cat_col}", model_name)
                if fig:
                    st.pyplot(fig)

            rendered_any = True
        except Exception as e:
            st.warning(f"Chart rendering skipped: {e}")

    # === Handle references / links cleanly ===
    refs = re.findall(r"(https?://\S+|ftp://\S+)", ai_output)
    if refs:
        with st.expander(" References & Sources"):
            for r in refs:
                st.markdown(f"- [{r}]({r})")

    if not rendered_any:
        st.warning("No tables, inline percentages, numbered lists, or JSON charts detected.")

# --- Define export_report first ---
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
import json
import os
from io import BytesIO


def export_report(result):
    """
    Accepts str | dict | list | tuple and safely builds a PDF in memory.
    Returns PDF bytes compatible with Streamlit download.
    """

    def to_text(obj):
        if isinstance(obj, tuple):
            return to_text(obj[1]) if len(obj) > 1 else to_text(obj[0])
        if isinstance(obj, dict):
            preferred_order = [
                "combined", "summary", "text",
                "gemma", "qwen", "llama", "openrouter",
                "raw_output", "tables", "charts", "meta"
            ]
            parts = []
            for k in preferred_order:
                if k in obj:
                    parts.append(f"## {k.upper()}\n{to_text(obj[k])}")
            for k, v in obj.items():
                if k not in preferred_order:
                    parts.append(f"## {str(k).upper()}\n{to_text(v)}")
            return "\n\n".join(parts) if parts else json.dumps(obj, indent=2, ensure_ascii=False)
        if isinstance(obj, list):
            return "\n".join(f"- {to_text(item)}" for item in obj)
        return str(obj)

    # Normalize any input to string
    text = to_text(result)
    if not isinstance(text, str):
        text = json.dumps(text, indent=2, ensure_ascii=False)

    # Build PDF in memory
    pdf = FPDF()
    pdf.add_page()
    try:
        font_path = os.path.join(os.path.dirname(__file__), "DejaVuSans.ttf")
        pdf.add_font('DejaVu', "", font_path, uni=True)
        pdf.set_font('DejaVu', size=12)
    except Exception:
        pdf.set_font('Arial', size=12)

    pdf.multi_cell(0, 8, txt=text)

    # ok fpdf returns bytearray → convert to bytes for Streamlit
    return bytes(pdf.output(dest='S'))

# --- Streamlit UI: show buttons only post-analysis ---
if "result" in st.session_state and st.session_state.result:
    result = st.session_state.result
    # --- PDF Export ---
    pdf_file = export_report(st.session_state.result)
    st.download_button(
        label="⬇️ Download Harmony Report (PDF)",
        data=pdf_file,
        file_name="Harmony_report.pdf",
        mime="application/pdf"
    )

    # --- Word Export ---
    doc = Document()
    doc.add_heading("Harmony Partner AI Report", 0)
    
    # --- Global font settings ---
    style = doc.styles["Normal"] 
    font = style.font
    font.name = "DejaVu Sans"
    font.size = Pt(11)
 
    # --- Body content ---
    doc.add_paragraph(str(st.session_state["result"]))
    
    # ✅ Save to memory instead of disk
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # ✅ Direct download (no file permission issues ever)
    st.download_button(
  	label="⬇️ Download Harmony Report (Word)",
  	data=buffer,
        file_name="Harmony_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

# --- Post-Analysis UI ---
if "result" in st.session_state and st.session_state.result:
    # 1. Always generate the latest report bytes from session state
    report_pdf_bytes = bytes(export_report(st.session_state.result))

    # --- Standard Download Buttons ---
    st.download_button("⬇️ Download PDF", data=report_pdf_bytes, file_name="Harmony_Report.pdf")

    # --- Email Logic ---
    if st.button("📧 Send Report via Email"):
        # Fetch recipients from your config
        raw_recipients = config.get("recipients", [])
        recipients = [e for e in raw_recipients if is_valid_email(e)]

        if not recipients:
            st.error("No valid email recipients found in configuration.")
        else:
            with st.spinner("Sending report to partners..."):
                success_count = 0
                for email in recipients:
                    # Directly passing the bytes from session state
                    success = send_email(
                        to_email=email,
                        subject="📊 Harmony AI Analysis Report",
                        body="Hello, your Harmony analysis report is attached as a PDF.",
                        attachment=report_pdf_bytes,
                        filename="Harmony_Report.pdf"
                    )
                    if success:
                        success_count += 1
                
                if success_count > 0:
                    st.success(f"Report successfully sent to {success_count} recipient(s).")

def log_crash(error_msg):
    crash_time = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    crash_path = LOG_DIR / f"crash_log_{crash_time}.txt"
    with open(crash_path, "w", encoding="utf-8") as f:
        f.write(f"[ERROR @ {crash_time}]\n{error_msg}")

def log_chat(user_id, model_used, model_selected, message):
    with open(CHAT_LOG_FILE, "a", encoding="utf-8") as f:
        now = datetime.datetime.now()
        f.write(f"{now} | {user_id} | Model: {model_used} | Method: {model_selected} | Message: {message.strip()} | Time: {now.time()} | Success\n")


    
# === ⚡ Flashmind Groq Analyzer (Resilient + Unified Fallback Edition) ===
def analyze_with_groq(prompt):
    import streamlit as st
    import time

    try:
        headers = {
            "Authorization": f"Bearer {api_key_column.get('groq_key', '')}",
            "Content-Type": "application/json"
        }

        if not headers["Authorization"]:
            st.warning("🔑 Flashmind API key missing. Please configure it first.")
            return {
                "Analysis 1": "Missing API key",
                "Analysis 2": "",
                "Summary": ""
            }

        # --- Safe model call with retry ---
        def call_groq_model(model_name, prompt, timeout=60, retries=3):
            """Primary Groq model caller with retry logic (no fallback here)."""
            for attempt in range(retries):
                try:
                    res = requests.post(
                        "https://api.groq.com/openai/v1/chat/completions",
                        headers=headers,
                        json={
                            "model": model_name,
                            "messages": [{"role": "user", "content": prompt}]
                        },
                        timeout=timeout,
                    )
                    data = res.json()

                    if "choices" in data and data["choices"]:
                        return data["choices"][0]["message"]["content"].strip()
                    elif "error" in data and "rate_limit" in str(data["error"]).lower():
                        wait_time = 20 * (attempt + 1)
                        st.info(f"⚙️ Optimized engine generating response... waiting {wait_time}s.")
                        time.sleep(wait_time)
                        continue
                    else:
                        raise ValueError(f"Invalid Groq output: {data}")

                except Exception as e:
                    if attempt < retries - 1:
                        st.warning(f"Retry {attempt + 1}/{retries} due to: {e}")
                        time.sleep(5)
                        continue
                    else:
                        st.error(f"Model {model_name} failed after all retries.")
                        return None
        
        # ============================================================
        # ✅ Unified fallback handler using ONLY TWO MODELS
        # ============================================================
        FALLBACK_MODELS = [
            "llama-3.1-8b-instant",
            "meta-llama/llama-prompt-guard-2-86m",
            "openai/gpt-oss-20b",
            "openai/gpt-oss-120b"
        ]

        # --- Fallback handler (triggered only at the end) ---
        def call_fallback_model(prompt):
            """Final rescue fallback: OpenAI GPT-OSS-20B."""
            try:
                res = requests.post(
                    "https://api.groq.com/openai/v1/chat/completions",
                    headers=headers,
                    json={
                        "model": FALLBACK_MODELS,
                        "messages": [{"role": "user", "content": prompt}]
                    },
                    timeout=120,
                )
                data = res.json()
                if "choices" in data and data["choices"]:
                    st.success("✅ Fallback engine recovered successfully.")
                    return data["choices"][0]["message"]["content"].strip()
                else:
                    st.error("Fallback model returned empty output.")
                    return "All engines failed to produce output."
            except Exception as e:
                st.error(f"❌ Fallback model also failed: {e}")
                return "Error: All models failed to generate output."

        # --- Step 1: Run Analysis 1 ---
        st.write("🧠 Running Analysis 1 ..")
        out1 = call_groq_model("groq/compound-mini", prompt, timeout=90)

        # --- Step 2: Run Analysis 2 ---
        st.write("🧩 Running Analysis 2 ...")
        out2 = call_groq_model("llama-3.1-8b-instant", prompt, timeout=90)

        # --- Step 3: Create blended summary ---
        blend_prompt = f"""Here are two analyses of the same market data:

Analysis 1:
{out1 or '[Analysis 1 failed]'}

Analysis 2:
{out2 or '[Analysis 2 failed]'}

1. SYNTHESIZED INSIGHTS  
   - Integrate findings from Analysis 1 and Analysis 2  
   - Highlight converging and diverging observations  

2. RECOMMENDED ACTIONS  
   - Clearly prioritized actions  
   - Practical and implementable steps  

3. 2026-ALIGNED STRATEGIC VIEW  
   - Key trends relevant to 2026  
   - Risks, opportunities, and strategic implications  

4. USER VS DOER PERSPECTIVES  
   - Explicit comparison of viewpoints  
   - Areas of alignment and conflict  

5. SUGGESTED VISUALS  
   - Recommended charts (bar, pie, trend)  
   - Tables for comparisons or contributions  

6. EXECUTIVE SUMMARY  
   - Professional, structured summary  
   - Suitable for Board, audit, or leadership review  

OUTPUT RULES:
- Use numbered headings and clear sections  
- Short paragraphs with spacing (maximum 3 lines)  
- Use tables where comparison adds clarity  
- Neutral, analyst / CFO briefing tone  
- No emojis or casual language  
- High clarity with audit-safe wording."""
        
        st.caption("Synthesizing optimized Flashmind summary...")
        summary = call_groq_model("groq/compound", blend_prompt, timeout=90)

        # --- Step 4: Final fallback decision ---
        if not out1 and not out2 and not summary:
            st.warning("⚠️ All Groq analyses failed. Switching to fallback engine...")
            final_output = call_fallback_model(prompt)
            return {
                "Analysis 1": "Fallback engine used.",
                "Analysis 2": "",
                "Summary": final_output
            }

        if not summary:
            st.warning("Summary failed — generating fallback summary...")
            summary = call_fallback_model(blend_prompt)

        # --- Step 5: Return final result ---
        return {
            "Analysis 1": (out1 or "Failed to generate Analysis 1").strip(),
            "Analysis 2": (out2 or "Failed to generate Analysis 2").strip(),
            "Summary": summary.strip()
        }

    except Exception as e:
        st.error(f"❌ Flashmind Analysis server issue: {e}")
        # Last-resort fallback
        return {
            "Analysis 1": "Error generating Analysis 1",
            "Analysis 2": "Error generating Analysis 2",
            "Summary": call_fallback_model(prompt)
        }
        
ANALYSIS_FALLBACK_MODELS = [ 
    "openai/gpt-oss-20b:free",
    "deepseek/deepseek-r1-distill-llama-70b:free",
    "meta-llama/llama-3.2-3b-instruct:free",
    "nvidia/nemotron-nano-12b-v2-vl:free",
    "nvidia/nemotron-nano-9b-v2:free",
    "x-ai/grok-4.1-fast:free"
]

def analyze_with_openrouter(prompt, stream=False): 
    import time, streamlit as st

    try:
        headers = {
            "Authorization": f"Bearer {api_key_column['openrouter']}",
            "Content-Type": "application/json"
        }

        # === Helper: Safe model call with fallback ===
        def call_openrouter(model_name, prompt, stream=False, timeout=60, retries=2):
            url = "https://openrouter.ai/api/v1/chat/completions"
            payload = {
                "model": model_name,
                "messages": [{"role": "user", "content": prompt}],
                "stream": stream
            }

            for attempt in range(retries):
                try:
                    res = requests.post(url, headers=headers, json=payload, timeout=timeout)

                    if res.status_code != 200 or not res.text:
                        raise ValueError(f"Invalid response: {res.status_code}, {res.text[:100]}")

                    try:
                        data = res.json()
                        if "choices" not in data or not data["choices"]:
                            raise ValueError("No choices in response.")
                        return data["choices"][0]["message"]["content"].strip()
                    except Exception:
                        return res.text.strip()

                except Exception as e:
                    if attempt < retries - 1:
                        st.info(f"⚙️ Optimized Omnicore engine generating response (retry {attempt + 1})...")
                        time.sleep(8 * (attempt + 1))
                    else:
                        st.warning(f" Primary engine failed: {e}. Switching to fallback models.")
                        # Fallback through list of models
                        for fallback_model in ANALYSIS_FALLBACK_MODELS:
                            try:
                                fallback_payload = {
                                    "model": fallback_model,
                                    "messages": [{"role": "user", "content": prompt}],
                                    "stream": stream
                                }
                                fallback_res = requests.post(url, headers=headers, json=fallback_payload, timeout=timeout)
                                fallback_data = fallback_res.json()
                                if "choices" in fallback_data and fallback_data["choices"]:
                                    return fallback_data["choices"][0]["message"]["content"].strip()
                            except Exception as e2:
                                st.info(f"⚙️ Fallback model {fallback_model} failed: {e2}")
                        raise RuntimeError("All fallback models failed.")

        # === Step 1: Run Analysis 1 ===
        st.write("🧠 Optimized engine-Omnicore generating Analysis 1...")
        out1 = call_openrouter("google/gemma-3-4b-it", prompt, stream=stream)

        # === Step 2: Run Analysis 2 ===
        st.write("🧩 Optimized engine-Omnicore generating Analysis 2...")
        out2 = call_openrouter("deepseek/deepseek-r1-distill-llama-70b", prompt, stream=stream)

        # === Step 3: Combine Outputs ===
        blend_prompt = f"""Here are two AI responses to the same input prompt.

Analysis 1:
{out1}

Analysis 2:
{out2}

Generate a final, consolidated summary suitable for executive, Board, and audit review.

The summary must include the following sections:

1. KEY INSIGHTS  
   - Critical observations derived from the analysis  
   - Supporting data points and examples relevant to 2024–2026  

2. RECOMMENDED ACTION POINTS  
   - Clearly defined, practical, and prioritized actions  
   - Focus on feasibility and impact  

3. COMPARATIVE ANALYSIS  
   - A structured comparison table for clarity  
   - Highlight differences, trade-offs, or alignment across options  

4. USER VS DOER PERSPECTIVES  
   - Distinct viewpoints from both User and Doer  
   - Areas of convergence and divergence  

5. REPORT-READY STRUCTURED TEXT  
   - Professionally written narrative suitable for reports  
   - No model identifiers, system credits, or meta commentary  

OUTPUT RULES:
- Use numbered headings and clearly separated sections  
- Keep paragraphs concise (maximum 3 lines each)  
- Use tables where comparison improves clarity  
- Maintain a neutral Analyst / CFO briefing tone  
- No emojis or casual language  
- High clarity with audit-safe wording."""

        st.write("🧩 Synthesizing final optimized summary...")
        combined_summary = call_openrouter("openai/gpt-oss-20b:free", blend_prompt, stream=stream, timeout=90)

        # === Step 4: Display ===
        full_output = f"""**Analysis 1:**\n{out1}\n\n
**Analysis 2:**\n{out2}\n\n
**Final Summary:**\n{combined_summary}"""

        st.subheader("📊 Omnicore Market Analysis Results")
        st.markdown(full_output)
        st.success("✅ Optimized engine completed all analysis successfully.")

        # Save session result
        st.session_state.result = {
            "Analysis 1": out1,
            "Analysis 2": out2,
            "Summary": combined_summary
        }

        # === Step 5: Optional Chart Rendering ===
        try:
            render_charts_from_ai_output(combined_summary)
        except Exception as e:
            st.warning(f"Chart rendering skipped: {e}")

        return combined_summary.strip()

    except Exception as e:
        st.error(f"❌ Omnicore Analyzer failed: {e}")
        st.info("⚙️ Attempting fallback engines...")

        # Try all fallback models
        for fallback_model in ANALYSIS_FALLBACK_MODELS:
            try:
                return call_openrouter(fallback_model, prompt, stream=stream)
            except Exception as e2:
                st.info(f"⚙️ Fallback model {fallback_model} failed: {e2}")
        st.error("All fallback models failed.")
        return "Analysis failed. Please retry later."


import streamlit as st


# === Groq Market Analysis Analyzer ===
def analyze_with_groq_market_analysis(m_prompt):
    """
    Generates a detailed Market Analysis report using Flashmind models.
    Optimized to generate a single robust report using the primary model, falling back
    to others only if necessary. 
    """

    # --- Get API Key ---
    api_keys = globals().get('api_key_column', st.session_state.get('api_key_column', {}))
    groq_market_key = api_keys.get('groq_market', "")
    
    if not groq_market_key:
        st.warning("⚠️ Flashmind API key is missing. Please enter your Flashmind API key to proceed.")
        return "Missing Flashmind Market API key."

    headers = {
        "Authorization": f"Bearer {groq_market_key}",
        "Content-Type": "application/json"
    }

    # --- STRICTLY USER DEFINED MODELS ---
    # Optimized Strategy: Use gpt-oss-120b first. If it works, stop. If not, try next.
    groq_models = [
        "openai/gpt-oss-120b",   # Primary (Fast & Effective)
        "groq/compound-mini",     # Secondary
        "openai/gpt-oss-20b",     # Fallback 1
        "llama-3.1-8b-instant"     # Fallback 2
    ]
    
    responses = []

    def call_groq_model(model_name, prompt, timeout=60):
        try:
            res = requests.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers=headers,
                json={"model": model_name, "messages": [{"role": "user", "content": prompt}]},
                timeout=timeout
            )
            if res.status_code != 200:
                return None # Trigger fallback
                
            data = res.json()
            if "choices" not in data or not data["choices"]:
                return None
            return data["choices"][0]["message"]["content"]
        except Exception:
            return None

    try:
        st.info("📊 Running Flashmind Market Analysis...")

        # --- Step 1: Run Primary Model (with Fallback) ---
        # Optimization: We loop through models but BREAK after the first success.
        model_used_label = ""
        
        for idx, model in enumerate(groq_models, start=1):
            
            # Generic display name to hide backend details
            display_name = f"Flashmind Model {idx}"
            status_msg = f"Optimizing {display_name}..."

            with st.spinner(status_msg):
                response_text = call_groq_model(model, m_prompt)
                
                if response_text:
                    # Success! Store result and STOP trying other models.
                    responses.append((display_name, response_text.strip()))
                    model_used_label = display_name
                    break  # <--- CRITICAL: Ensures only ONE model report is generated
                else:
                    # Silently continue to next model if this one fails
                    continue

        # --- Step 2: Synthesis / Formatting ---
        if responses:
            # Prepare the single report for formatting
            combined_text = "\n\n".join([f"=== {label} ===\n{text}" for label, text in responses])
            
            blend_prompt = f"""Here is a raw Market Analysis:

{combined_text}

TASK OBJECTIVE

Refine the provided inputs into a **Final Consolidated Market Analysis Summary**
suitable for executive, Board, and audit review.

DELIVERABLE REQUIREMENTS

1. STRUCTURED CONSOLIDATION  
   - Synthesize all provided analyses into a single, coherent narrative  
   - Highlight the most material insights and conclusions  

2. ACTIONABLE RECOMMENDATIONS  
   - Present clearly prioritized and implementable recommendations  
   - Link each recommendation to observed market insights or data  

3. PERSPECTIVE SEPARATION  
   - Distinctly separate and label:
     a. User Perspective  
     b. Doer Perspective  
   - Clearly explain differences, trade-offs, and areas of alignment  

4. VISUAL REPRESENTATION (TEXT-BASED)  
   - Include Bar or Pie chart representations using:
     - Markdown data tables, or  
     - Clear textual chart descriptions suitable for rendering  
   - Ensure numeric or percentage values are provided where charts are suggested  

5. AUTHORITATIVE CONTEXT (2026)  
   - Reference credible, authoritative insights relevant to the year 2026  
   - Contextualize references to support conclusions and recommendations  

OUTPUT RULES

- Use numbered headings and clearly defined sections  
- Keep paragraphs concise (maximum 3 lines each)  
- Use tables where comparisons or distributions improve clarity  
- Maintain a neutral Analyst / CFO briefing tone  
- No emojis, no casual language, no model identifiers  
- Use precise, audit-safe wording suitable for regulatory and leadership review  
"""
            
            final_summary = "Summary generation failed."
            
            # --- Synthesizing Summary ---
            with st.spinner("Synthesizing Flashmind Market Report..."):
                # Use the same priority list for the summarizer
                for model in groq_models:
                    temp_summary = call_groq_model(model, blend_prompt)
                    if temp_summary:
                        final_summary = temp_summary
                        break

            # --- Step 3: Fetch References ---
            refs_md = ""
            try:
                if 'get_references' in globals():
                    refs = get_references(m_prompt, online_mode=True)
                    if refs:
                        for r in refs:
                            if isinstance(r, dict) and "url" in r and "title" in r:
                                refs_md += f"- [{r['title']}]({r['url']})\n"
                            else:
                                refs_md += f"- {r}\n"
                    else:
                        refs_md = "No specific references found."
                else:
                    refs_md = "Reference module not loaded."
            except Exception:
                refs_md = "References unavailable."

            # --- Step 4: Combine all in one box ---
            full_report = f"""
# 📌 Analysis Report

{combined_text}

# 📄 Final Consolidated Market Analysis Summary

{final_summary}

# 🔗 References

{refs_md}
"""

            st.subheader("📄 Market Analysis Report")
            st.text_area(
                label="Complete Report",
                value=full_report,
                height=800,
                max_chars=None,
                key="market_analysis_report"
            )

            # --- Step 5: Render charts ---
            try:
                if 'render_charts_from_ai_output' in globals():
                    render_charts_from_ai_output(final_summary)
            except Exception:
                pass

            return full_report.strip()

        return "All Flashmind analyses failed."

    except Exception as e:
        st.error(f"Flashmind Analysis failed: {e}")
        return "Failed to generate Market Analysis."
    

import time
import streamlit as st

# =====================================================
# INTERNAL OMNICORE FALLBACK MODELS (HIDDEN)
# =====================================================
ANALYSIS_FALLBACK_MODELS = [
    "openai/gpt-oss-20b:free",
    "qwen/qwen3-next-80b-a3b-instruct:free",
    "deepseek/deepseek-r1-distill-llama-70b:free",
    "meta-llama/llama-3.2-3b-instruct:free",
    "nvidia/nemotron-nano-12b-v2-vl:free",
    "nvidia/nemotron-nano-9b-v2:free"
]

# =====================================================
# BUILD ENGINEERING / SCIENCE PROMPT
# =====================================================
def build_o_prompt(user_comment, doer_comment, final_input, focus="science_research"):
    """
    Builds a single Omnicore research prompt for OpenAI GPT engine.
    """
    return f"""
### Science, Research & Engineering

User Input:
{final_input}

User Context:
{user_comment}

Doer Context:
{doer_comment}

Focus Area:
{focus}

You are a technical analysis engine.

Provide a full, detailed engineering, science or research analysis (2026+), covering: 
- Overview and purpose
- Types and classifications
- Scientific principles and theory
- Functional mechanisms
- Depth analysis with examples
- Comparison with alternatives
- Recommendations and best practices
- Research trends and future developments
- Standards, technical papers, and credible references

Rules:
- Provide only objective, evidence-based, and technical explanations.
- Use scientific consensus, industry standards, and empirical data.
- No emotional language, personal opinions, or moral judgments.
- No medical, emotional, or legal advice.
- High-level factual info only for overlapping domains.
- Never simulate emotions or intentions.
- Do not invent policies or disclaimers unless required.

Output format:
- Use numbered headings and clearly defined sections  
- Keep paragraphs concise (maximum 3 lines each)  
- Use tables where comparisons or distributions improve clarity  
- Maintain a neutral Analyst / CFO briefing tone  
- No emojis, no casual language, no model identifiers  
- Use precise, audit-safe wording suitable for regulatory and leadership review  

User Query:
{final_input}
"""

# =====================================================
# ANALYZE PROMPT USING OPENAI / OMNICORE ENGINES
# =====================================================
def analyze_with_openai(o_prompt):
    """
    Runs optimized Omnicore engines via OpenRouter.
    Retries and fallback handled internally.
    Returns output text only.
    """
    headers = {
        "Authorization": f"Bearer {api_key_column['openrouter']}",
        "Content-Type": "application/json",
        "HTTP-Referer": "http://localhost",
        "X-Title": "Engineering Analyzer"
    }

    retries = 2
    timeout = 120

    for model in ANALYSIS_FALLBACK_MODELS:
        for attempt in range(retries):
            try:
                resp = requests.post(
                    "https://openrouter.ai/api/v1/chat/completions",
                    headers=headers,
                    json={
                        "model": model,
                        "messages": [{"role": "user", "content": o_prompt}],
                        "max_tokens": 2500
                    },
                    timeout=timeout
                )

                if resp.status_code != 200:
                    raise RuntimeError(f"Engine response error: {resp.status_code}")

                data = resp.json()
                output = "".join(
                    c.get("message", {}).get("content", "")
                    for c in data.get("choices", [])
                ).strip()

                if output:
                    # Optional chart rendering
                    try:
                        render_charts_from_ai_output(output)
                    except Exception:
                        st.warning("Chart rendering skipped.")
                    return output

            except Exception:
                if attempt < retries - 1:
                    st.warning("⏳ Retrying optimized Omnicore engines…")
                    time.sleep(5)
                else:
                    break

    st.error("⚠️ Optimized Omnicore engines unavailable.")
    return "Omnicore analysis failed."


# =====================================================
# USAGE EXAMPLE
# =====================================================
# o_prompt = build_o_prompt(user_comment, doer_comment, final_input)
# result = analyze_with_openai(o_prompt)

    
ANALYSIS_FALLBACK_MODELS = [ 
    "openai/gpt-oss-20b:free",
    "openai/gpt-oss-120b:free",
    "deepseek/deepseek-r1-distill-llama-70b:free",
    "meta-llama/llama-3.2-3b-instruct:free",
    "qwen/qwen3-next-80b-a3b-instruct:free",
    "google/gemma-4-31b-it:free",
    "x-ai/grok-4.1-fast:free"
]

def build_tool_prompt(user_comment, doer_comment, final_input, focus="engineering_tool_design"):
    """
    Builds a unified engineering + tool design + manufacturing prompt.
    Matches style of build_market_prompt().
    """

    return f"""
### Engineering & Tool Design Analysis

User Input:
{final_input}

User Context:
{user_comment}

Doer Context:
{doer_comment}

Focus Area: {focus}

Provide a full detailed analysis latest 2026 including:
- Tool design and Pattern Making / Tooling & strategy
- Raw Material & Casting Machine selection, Moulding Section, machinery requirements, capacities
- Melting & Pouring, composition, Shakeout + Fettling, strengths, properties
- Heat Treatment Section, CNC Machining Line Setup with pre and finish machining
- Balancing Cell, FFT / NVH Validation Lab, Surface Treatment Line, Inspection & QC Section, Packing & Storage
- Categories, types, variants, classifications
- Working principle and operating mechanism
- Setup requirements, cost insights, process optimization, Safety, durability and performance considerations
- Sales Strategies, Customer requirements and marketing related information in detail 
- Recommended Plant Layout (Area Wise), Manpower Requirement, Industrial setup, plant layout, production flow
- Technologies involved, advanced methods, automation options with recommended companies. 
- Competition, market comparisons, innovation opportunities and names of companies. 


Return the entire analysis in a clean, structured engineering format.

OUTPUT RULES
- Use numbered headings and clearly defined sections  
- Keep paragraphs concise (maximum 3 lines each)  
- Use tables where comparisons or distributions improve clarity  
- Maintain a neutral Analyst / CFO briefing tone  
- No emojis, no casual language, no model identifiers  
- Use precise, audit-safe wording suitable for regulatory and leadership review  
"""

def analyze_with_gemini(tool_prompt):
    """
    Gemini analysis using OpenRouter with automatic fallback models.
    Uses the SAME retry + fallback structure as analyze_with_openai_deepseek().
    """

    import streamlit as st
    import time

    try:
        # ------------ HEADERS ------------
        OPENROUTER_HEADERS = {
            "Authorization": f"Bearer {api_key_column['openrouter']}",
            "Content-Type": "application/json",
            "HTTP-Referer": "http://localhost",
            "X-Title": "Engineering Analyzer"
        }

        # ------------ PROMPT ------------
        gemini_prompt = (
            "Provide detailed engineering, tooling, industrial and factory analysis in detail, examples and recommendations:\n"
            "- Materials, structure, composition\n"
            "- Tool design, machining, tolerances, fixtures\n"
            "- Production workflow, plant layout\n"
            "- Safety, competitors, optimization, cost\n\n"
            f"Tool Prompt: {tool_prompt}"
        )

        # ========== FALLBACK MODELS ==========
        ANALYSIS_FALLBACK_MODELS = [
            "openai/gpt-oss-120b:free",              # primary
            "deepseek/deepseek-r1-distill-llama-70b:free",
            "openai/gpt-oss-20b:free",
            "nvidia/nemotron-nano-12b-v2-vl:free",
            "google/gemma-4-31b-it:free",
            "x-ai/grok-4.1-fast:free"
        ]

        # ------------ REQUEST FUNCTION ------------
        def call_openrouter(model_name, prompt_text, timeout=120, retries=2):
            for attempt in range(retries):
                try:
                    # ⭐ YOUR EXACT POST BLOCK ⭐
                    resp = requests.post(
                        "https://openrouter.ai/api/v1/chat/completions",
                        headers=OPENROUTER_HEADERS,
                        json={
                            "model": model_name,
                            "messages": [{"role": "user", "content": prompt_text}],
                            "max_tokens": 3300
                        },
                        timeout=timeout
                    )
                    # ---------------------------

                    if resp.status_code != 200:
                        raise ValueError(f"HTTP {resp.status_code}: {resp.text[:200]}")

                    data = resp.json()
                    output = "".join(
                        c.get("message", {}).get("content", "") + "\n"
                        for c in data.get("choices", [])
                    ).strip()

                    if not output:
                        raise ValueError("Empty response from Omnicore")

                    return output

                except Exception as e:
                    if attempt < retries - 1:
                        st.warning(f"Retry {attempt+1}/{retries} for {model_name}: {e}")
                        time.sleep(5)
                    else:
                        st.warning(f"Model {model_name} failed: {e}")
                        return None

        # ========== RUN THROUGH FALLBACK MODELS ==========
        final_output = None
        used_model = None

        for model in ANALYSIS_FALLBACK_MODELS:
            final_output = call_openrouter(model, gemini_prompt)
            if final_output:
                used_model = model
                break

        if not final_output:
            st.error("⚠ All models (Gemini + fallback models) failed.")
            return "Gemini analysis failed."

        # ------------ SAVE & RENDER ------------
        st.session_state.result = {
            "Omnicore": final_output,
            "Used_Model": used_model
        }

        try:
            render_charts_from_ai_output(final_output)
        except Exception as e:
            st.warning(f"Chart rendering skipped: {e}")

        return final_output

    except Exception as e:
        st.error(f"Gemini engine crashed: {e}")
        return "Gemini crashed."


def analyze_with_openai_deepseek(prompt):
    """
    Analyze using Omnicore with full fallback logic.
    Updated with SAME working structure used in Gemini engine.
    """

    import streamlit as st
    import time

    try:
        # ------------ HEADERS ------------
        OPENROUTER_HEADERS = {
            "Authorization": f"Bearer {api_key_column['openrouter_fin']}",
            "Content-Type": "application/json",
            "HTTP-Referer": "http://localhost",
            "X-Title": "Engineering Analyzer"
        }

        # ------------ ENGINEERING PROMPT ------------
        engineering_prompt = f"""
        TASK OBJECTIVE
 
        Analyze the given query and provide **detailed, professional engineering insights**
        aligned with current industry practices and developments relevant to **2026**.

        DELIVERABLE REQUIREMENTS
        1. TECHNICAL ANALYSIS
            Components and Engineering Composition
                        Complete breakdown of all major subsystems and assemblies
                        Material selection and engineering justification (aluminum alloys, polymers, steel grades, composites, coatings, structural materials, etc.)
                        Mechanical, electrical, and structural design architecture
                        Thermal, structural, vibration, and load-bearing characteristics
                        Engineering constraints affecting performance, cost, and manufacturability
            Design intent and functional role of key components
                        Functional purpose of each subsystem
                        Failure-prone areas and engineering vulnerabilities
                        Critical tolerances and engineering sensitivities
                        Heat dissipation, structural integrity, and durability considerations

        2. MANUFACTURING & PROCESS ENGINEERING
            Machining methods, tooling, and tolerances
                        CNC machining, die casting, injection molding, stamping, extrusion, fabrication, laser cutting, robotic welding (as applicable)
                        Tooling requirements, dies, molds, fixtures, and jigs
                        Critical tolerances and dimensional stability requirements
            Material engineering characteristics
                        Thermal conductivity considerations
                        Strength and durability optimization
                        Corrosion resistance and environmental protection
                        Weight optimization vs strength trade-offs
                        Cost vs performance trade-offs in material selection
            Manufacturing processes and process controls
                        Raw material sourcing (India and global supply chain options)
                        Manufacturing process flow from raw material to finished product
                        Surface finishing (anodizing, coating, powder coating, plating, polishing)
                        Assembly processes (manual, semi-automated, automated assembly lines)
                        Recommended production line configuration
            Testing and quality inspection stages
                        Incoming material inspection
                        In-process inspection
                        Final product inspection
                        Reliability and lifecycle testing
            Machinery and factory setup recommendations (MANDATORY)
                        Required machinery list with purpose
                        Approximate machinery cost in INR
                        Recommended Indian and global suppliers
                        Automation vs manual production comparison
            Scalability and manufacturability considerations
                        Production scalability from small to large scale
                        Automation opportunities
                        Bottlenecks and manufacturing constraints
                        Production capacity planning

          3. PRODUCTIVITY & OPTIMIZATION
            Equipment efficiency and utilization
                        Machine utilization optimization
                        Throughput improvement opportunities
                        Labor vs automation optimization
            Process improvements and cost optimization opportunities
                        Manufacturing inefficiencies and engineering root causes
                        Engineering solutions to reduce cost
                        Alternative materials or manufacturing methods to reduce cost
                        Energy efficiency optimization
                        Automation ROI analysis
            Yield improvement and waste reduction approaches
                        Scrap generation causes
                        Defect sources and failure modes
                        Engineering solutions to improve yield
                        Waste reduction and efficiency improvement strategies
            Factory setup cost breakdown (MANDATORY, INR-based)
                        Machinery cost
                        Tooling cost
                        Factory infrastructure cost
                        Electrical and utilities setup cost
                        Labor setup cost
                        Testing equipment cost
                        Total factory setup cost (Small scale, Medium scale, Large scale)
            Per unit manufacturing cost breakdown (INR-based)

          4. PERFORMANCE, SAFETY & RELIABILITY
            Performance benchmarks and operational limits
                        Structural performance
                        Thermal limits
                        Operational efficiency limits
                        Environmental performance limits
            Safety considerations and compliance aspects
                        BIS standards (India)
                        AIS standards (automotive if applicable)
                        ISO and industrial compliance standards
                        Environmental and safety compliance
            Reliability, durability, and lifecycle expectations
                        Expected lifespan
                        Failure modes and root causes
                        Maintenance requirements
                        Durability under Indian environmental conditions (heat, dust, vibration)

           5. COMPETITIVE & INDUSTRY BENCHMARKING
            Comparison with competing technologies, technical trending latest ideas 2026 or solutions
                        Comparison with traditional manufacturing methods
                        Comparison with advanced automated manufacturing
                        Global vs Indian manufacturing cost comparison
                        Benchmark performance and manufacturing efficiency
            Identification of differentiators and gaps
                        Engineering advantages
                        Manufacturing efficiency advantages
                        Cost advantages or disadvantages
                        Technology gaps and improvement opportunities
            Latest industry trends (2026–2026)
                        Automation and smart manufacturing trends
                        Industry 4.0 integration
                        Robotics and AI-driven manufacturing trends
                        India manufacturing trends (PLI schemes, localization, Make in India)

          6. COMMERCIAL & MARKET CONTEXT (ENGINEERING-LINKED)
            Engineering-driven value propositions
                        Engineering advantages affecting cost, durability, and efficiency
                        Manufacturing advantages affecting competitiveness
            Market-facing recommendations (technical differentiators only)
                        Engineering improvements to improve market competitiveness
                        Cost reduction engineering strategies
                        Localization opportunities
            Factory setup investment vs ROI analysis
                        Estimated ROI timelines
                        Production volume vs profitability analysis
            Relevant company or industry examples (2026–2026)
                        Indian manufacturers
                        Global manufacturers
                        OEM and Tier-1 supplier manufacturing practices

           7. PROBLEMS & ENGINEERING SOLUTIONS (MANDATORY SECTION)
            Identify real manufacturing and engineering problems and provide structured solutions.
            Format:
                        Problem
                        Root Cause
                        Engineering Solution
                        Expected Improvement (cost reduction %, lifespan increase %, efficiency increase %, etc.)
            Include:
                        Manufacturing problems
                        Cost problems
                        Reliability problems
                        Scalability problems
                        Quality problems


        OUTPUT RULES 
        - Use structured headings and numbered sections   
        - Keep paragraphs concise (maximum 4 lines each with each points covered and topic related elaboration) 
        - All costs In INR and detailed bifurcation  
        - Use tables where comparisons add clarity  
        - Maintain a neutral, professional engineering tone  
        - No emojis, no casual language, no marketing hype  
        - Use precise, audit-safe terminology suitable for technical and leadership review  

ENGINEERING QUERY:
{prompt}
"""

        # ------------ FALLBACK LIST (same as Gemini) ------------
        ANALYSIS_FALLBACK_MODELS = [
            "qwen/qwen3-next-80b-a3b-instruct:free",
            "deepseek/deepseek-r1-distill-llama-70b:free",
            "openai/gpt-oss-120b:free",
            "nvidia/nemotron-nano-12b-v2-vl:free",
            "google/gemma-4-31b-it:free",
            "x-ai/grok-4.1-fast:free"
        ]

        # ------------ REQUEST FUNCTION (same as Gemini) ------------
        def call_openrouter(model_name, prompt_text, timeout=120, retries=2):

            for attempt in range(retries):
                try:
                    # ⭐ EXACT SAME POST BLOCK USED IN GEMINI ⭐
                    resp = requests.post(
                        "https://openrouter.ai/api/v1/chat/completions",
                        headers=OPENROUTER_HEADERS,
                        json={
                            "model": model_name,
                            "messages": [{"role": "user", "content": prompt_text}],
                            "max_tokens": 3000
                        },
                        timeout=timeout
                    )
                    # ---------------------------------------------------

                    if resp.status_code != 200:
                        raise ValueError(f"HTTP {resp.status_code}: {resp.text[:200]}")

                    data = resp.json()
                    output = "".join(
                        c.get("message", {}).get("content", "") + "\n"
                        for c in data.get("choices", [])
                    ).strip()

                    if not output:
                        raise ValueError("Empty response from OpenRouter.")

                    return output

                except Exception as e:
                    if attempt < retries - 1:
                        st.warning(f"⏳ Retry {attempt+1}/{retries} for {model_name}: {e}")
                        time.sleep(5)
                    else:
                        st.warning(f"Model {model_name} failed: {e}")
                        return None

        # ------------ RUN THROUGH FALLBACK MODELS ------------
        final_output = None
        used_model = None

        for model_name in ANALYSIS_FALLBACK_MODELS:
            final_output = call_openrouter(model_name, engineering_prompt)
            if final_output:
                used_model = model_name
                break

        if not final_output:
            st.error("⚠️ All online models failed to respond. Please retry later.")
            return "Analysis failed. All online models unavailable."

        # ------------ SAVE & RENDER ------------
        st.session_state.result = {
            "Omnicore": final_output,
            "Used_Model": used_model
        }

        try:
            render_charts_from_ai_output(final_output)
        except Exception as e:
            st.warning(f"Chart rendering skipped: {e}")

        return final_output

    except Exception as e:
        st.error(f"Unexpected error in DeepSeek analysis: {e}")
        return "Analysis failed. Please retry later."



def analyze_with_openrouter_fin(f_prompt, perspective="user"):
    """
    Analyze financial documents using Omnicore.
    Runs BOTH Optimized Models Finance online.
    Falls back to TWO recommended financial models if any primary model fails,
    and finally falls back to offline Qwen + TinyDolphin.
    """

    import streamlit as st

    perspective = str(perspective).lower()

    ONLINE_FINANCE_PROMPT_USER = """You are a professional financial analyst. Analyze the following balance sheet or income statement from a User perspective. Provide insights on overall financial health, key trends, and potential risks. Use headings and subheadings for clarity."""
    ONLINE_FINANCE_PROMPT_DOER = """You are a professional financial analyst. Analyze the following balance sheet or income statement from a Doer perspective. Provide actionable recommendations, operational improvements, and steps to optimize liquidity, profitability, and solvency. Use headings, subheadings, and simple tables."""

    financial_prompt = ONLINE_FINANCE_PROMPT_DOER if perspective == "doer" else ONLINE_FINANCE_PROMPT_USER
    input_text = financial_prompt + "\n\n" + f_prompt

    # --- Primary OpenRouter Models ---
    openrouter_models = {
        "Analysis 1": "openai/gpt-oss-20b:free",
        "Analysis 2": "mistralai/mistral-7b-instruct:free"
    }

    # --- Recommended fallback financial models ---
    fallback_models = [
        "nvidia/nemotron-nano-12b-v2-vl:free",
        "meta-llama/llama-3.2-3b-instruct:free"
    ]

    headers = {
        "Authorization": f"Bearer {api_key_column.get('openrouter_fin', '')}",
        "Content-Type": "application/json"
    }
    url = "https://openrouter.ai/api/v1/chat/completions"

    combined_response = ""

    if not headers["Authorization"].strip():
        st.warning(" Omnicore key not provided. Falling back to offline model...")
        return analyze_with_qwen(f_prompt, perspective)

    try:
        # --- Step 1: Run primary models ---
        for label, model_name in openrouter_models.items():
            try:
                payload = {
                    "model": model_name,
                    "messages": [{"role": "user", "content": input_text}],
                    "temperature": 0.7,
                    "max_tokens": 4000
                }
                res = requests.post(url, headers=headers, json=payload, timeout=200)
                data = res.json()
                output = data.get("choices", [{}])[0].get("message", {}).get("content", "")
                if not output.strip():
                    raise ValueError(f"{label} returned empty response.")
                combined_response += f"\n\n--- {label} RESPONSE ---\n{output.strip()}"

            except Exception as inner_e:
                st.warning(f"{label} failed: {inner_e}. Trying fallback financial model...")
                # --- Step 2: Fallback to first available financial model ---
                for fallback_model in fallback_models:
                    try:
                        fallback_payload = {
                            "model": fallback_model,
                            "messages": [{"role": "user", "content": input_text}],
                            "temperature": 0.7,
                            "max_tokens": 4000
                        }
                        fallback_res = requests.post(url, headers=headers, json=fallback_payload, timeout=200)
                        fallback_data = fallback_res.json()
                        fallback_output = fallback_data.get("choices", [{}])[0].get("message", {}).get("content", "")
                        if fallback_output.strip():
                            combined_response += f"\n\n--- {label} (Fallback) RESPONSE ---\n{fallback_output.strip()}"
                            break
                    except Exception as fb_e:
                        st.warning(f"Fallback model {fallback_model} also failed: {fb_e}")

        # --- Step 3: Chart rendering ---
        try:
            render_charts_from_ai_output(combined_response)
        except Exception as e:
            st.warning(f"Chart rendering skipped (finance): {e}")

        st.session_state.result = combined_response.strip()
        return combined_response.strip()

    except Exception as e:
        st.warning(f"Omnicore Analysis failed: {e}\n🧠 Falling back to offline models...")
        return analyze_with_qwen(f_prompt, perspective)
            
# === Streamlit UI trigger ===
if st.button("Render Charts from AI Output"):
    if "result" in st.session_state and st.session_state.result:
        render_charts_from_ai_output(st.session_state.result)
    else:
        st.warning("No AI result available. Run analysis first.")

# === Run Analysis Button ===
if st.button("🔍 Run Analysis"):
    final_input = manual_input or extracted_text
    if not final_input:
        st.error("No input provided.")
    else:
        st.info("Generating AI output...")
        start_time = time.time()

        prompt = build_prompt(user_comment, doer_comment, final_input, model_type, user_priority)
        # Build the final prompt
        f_prompt = build_f_prompt(user_comment, doer_comment, final_input, model_type, user_priority, perspective="user"   # or "doer"
        )

        # === Engine Dispatcher ===
        if engine_type == "Offline":
            result = analyze_with_offline(prompt, offline_choice)
        elif engine_type == "Qwen-fin":
            result = analyze_with_qwen(f_prompt, user_priority)    
        elif engine_type == "Offline RAG":
            result = analyze_with_offline_rag(prompt)
        elif engine_type == "Flashmind RAG":
            result = analyze_with_groq(prompt)
        elif engine_type == "Omnicore RAG":
            result = analyze_with_openrouter(prompt)
        elif engine_type == "Flashmind Market":
            m_prompt = build_market_prompt(
                user_comment,
                doer_comment,
                final_input,
                market_focus="Market Research & Brand marketing",   # strictly Market
                perspective="user"          # can toggle to "doer" if needed
            )
            result = analyze_with_groq_market_analysis(m_prompt)
        elif engine_type == "OpenAi-GPT":
            o_prompt = build_o_prompt(
                user_comment,
                doer_comment,
                final_input,
                focus="science_engineering_research"
            )
            result = analyze_with_openai(o_prompt)
        elif engine_type == "Product-R&D-Omni":
            # Build engineering + tool-design prompt
            tool_prompt = build_tool_prompt(
                user_comment,
                doer_comment,
                final_input,
                focus="engineering_tool_design"   # custom profile for tools
            )
            result = analyze_with_gemini(tool_prompt)
        elif engine_type == "Product-Optimizations-Omni":
            result = analyze_with_openai_deepseek(prompt)
        elif engine_type == "Omnicore_Finance":
            result = analyze_with_openrouter_fin(f_prompt, user_priority)    
        else:
            result = (" Invalid Engine", "Engine not recognized.")

        # === Save result to session ===
        st.session_state.result = result
        
        log_user_activity(
            st.session_state["user_id"],
            "analysis_run",
            {
                "engine": engine_type,
                "method": model_type
            }
        ) 
        # === Standardize AI output text for charts & display ===
        if isinstance(result, tuple):
            source, content = result
            if isinstance(content, dict):
                # Merge all outputs including individual models and combined summary
                ai_output_text = "\n\n".join([f"{k.upper()} RESPONSE:\n{v}" for k, v in content.items()])
            else:
                ai_output_text = str(content)
        elif isinstance(result, dict):
            ai_output_text = "\n\n".join([f"{k.upper()} RESPONSE:\n{v}" for k, v in result.items()])
        else:
            ai_output_text = str(result)

        # === Render charts from AI output ===
        render_charts_from_ai_output(ai_output_text)

        # === Scrollable output box ===
        bg_color = "#f8f9fa" if not st.session_state.dark_mode else "#0E1117"
        text_color = "#000000" if not st.session_state.dark_mode else "#FAFAFA"

        safe_text = (
            ai_output_text.replace("&", "&amp;")
                          .replace("<", "&lt;")
                          .replace(">", "&gt;")
                          .replace("\n", "<br>")
        )

        st.markdown(
            f"""
            <div style="max-height: 600px; overflow-y: auto; overflow-x: auto;
                        border: 1px solid #888; padding: 10px;
                        white-space: pre-wrap; word-wrap: break-word;
                        background-color: {bg_color}; color: {text_color}; font-family: monospace;">
                {safe_text}
            </div>
            """,
            unsafe_allow_html=True
        )

        # === Show time taken ===
        duration = round(time.time() - start_time, 2)
        st.info(f"⏱️ Time Taken: {duration} seconds")

        # === Export report safely with all responses ===
        export_report(result)

# ===== Footer =====
st.markdown(
    """
    <hr style="margin-top: 40px; margin-bottom: 10px;">
    <div style="text-align: center; font-size: 13px; color: gray;">
        © 2026 All Rights Reserved.<br>
        Harmony-GPT can make mistakes.<br>
        Powered by Retrieval-augmented generation (RAG)
        Designed under patent theory by owner. (Kindly read consents & readme before using) <br>
        ("Contact 📱 : whatsapp chat ✔ 8383995759")   
    </div>
    """,
    unsafe_allow_html=True
)
